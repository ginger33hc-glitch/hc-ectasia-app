"""Professional PDF/DOCX export guard: reference appendices, concise identity notice, and safe serialization."""
from io import BytesIO
import re
from threading import RLock
import reports

_orig_eye_metrics=reports._eye_metrics
_orig_tomo_rows=reports._tomography_rows

def _eye_metrics_no_morphology(eye,locale="en"):
    return [(k,v) for k,v in _orig_eye_metrics(eye,locale) if k.strip().lower() not in {"morphology category","morfoloji kategorisi"}]

def _tomography_rows_no_morphology(extracted,eye_id,locale="en"):
    return [(k,v) for k,v in _orig_tomo_rows(extracted,eye_id,locale) if k.strip().lower() not in {"morphology","morfoloji"}]

reports._eye_metrics=_eye_metrics_no_morphology
reports._tomography_rows=_tomography_rows_no_morphology

from reportlab.platypus import KeepTogether,Paragraph,Table,TableStyle,Spacer
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt

IDENTITY_HEADING="PATIENT IDENTITY NOT VERIFIED - SURGEON CONFIRMATION REQUIRED"
PDF_BAD=[["Final BAD-D","CER-AI interpretation / action"],["<=1.6","NORMAL"],[">1.6 to <2.60","SUSPICIOUS - REVIEW / NOT CLEARED"],[">=2.60","ABNORMAL CORNEA - DO NOT PROCEED"]]
TOPOGRAPHY=[["Category","What to look for","ERSS points"],*[list(row) for row in reports.RANDLEMAN_TOPOGRAPHY_REFERENCE]]
ACTIVE_ERSS=[list(row) for row in reports.RANDLEMAN_ACTIVE_ERSS_REFERENCE]

_orig_pdf=reports.build_pdf
_orig_docx=reports.build_docx
_EXPORT_LOCK=RLock()

def _current_version():
    return str(getattr(reports,"APP_VERSION","unknown"))

def _sync_pdf_version(story):
    version=_current_version()
    for i,item in enumerate(story):
        if isinstance(item,Paragraph) and "Software v" in getattr(item,"text",""):
            text=re.sub(r"Software v[^ <]+",f"Software v{version}",item.text)
            story[i]=Paragraph(text,item.style)

def _sync_docx_version(data):
    version=_current_version()
    document=Document(BytesIO(data))
    changed=False
    for paragraph in document.paragraphs:
        if "Software v" in paragraph.text:
            for run in paragraph.runs:
                if "Software v" in run.text:
                    run.text=re.sub(r"Software v\S+",f"Software v{version}",run.text)
                    changed=True
            if not changed:
                paragraph.text=re.sub(r"Software v\S+",f"Software v{version}",paragraph.text)
                changed=True
    if not changed:return data
    output=BytesIO();document.save(output);return output.getvalue()

def _remove_pdf_identity_details(story,identity_heading=IDENTITY_HEADING):
    """Keep the identity-verification heading but suppress verbose source-by-source bullets."""
    start=next((i for i,x in enumerate(story) if isinstance(x,Paragraph) and getattr(x,'text','')==identity_heading),None)
    if start is None:return
    end=start+1
    while end < len(story):
        item=story[end]
        if isinstance(item,Paragraph) and getattr(getattr(item,'style',None),'name','')=='Section':break
        end+=1
    del story[start+1:end]

def build_pdf(payload):
    # The legacy report builder requires a temporary ReportLab build hook. Serialize the
    # complete hook lifetime so concurrent web requests cannot observe each other's patch.
    with _EXPORT_LOCK:
        locale=reports.normalize_locale(payload.get("locale"))
        tr=lambda value:reports.translate_text(value,locale)
        regular_font=reports.PDF_UNICODE_REGULAR if locale=="tr" else "Helvetica"
        bold_font=reports.PDF_UNICODE_BOLD if locale=="tr" else "Helvetica-Bold"
        orig_build=reports.SimpleDocTemplate.build
        def patched_build(doc,story,*a,**kw):
            _sync_pdf_version(story)
            _remove_pdf_identity_details(story,tr(IDENTITY_HEADING))
            idx=next((i for i,x in enumerate(story) if isinstance(x,Paragraph) and getattr(x,'text','')==tr('Interpretation note')),len(story))
            styles=reports.getSampleStyleSheet()
            sec=reports.ParagraphStyle('HCRefSection',parent=styles['Heading2'],fontName=bold_font,fontSize=10.5,leading=13,textColor=reports._rl(reports.NAVY),spaceBefore=10,spaceAfter=5)
            tiny=reports.ParagraphStyle('HCRefTiny',parent=styles['BodyText'],fontName=regular_font,fontSize=7.2,leading=9,textColor=reports._rl(reports.GRAY))
            ref_head=reports.ParagraphStyle('HCRefHead',parent=tiny,fontName=bold_font,textColor=colors.white)
            ref_cell=reports.ParagraphStyle('HCRefCell',parent=tiny,fontName=regular_font,textColor=reports._rl(reports.INK))
            bad_rows=[[Paragraph(tr(cell),ref_head if row_index==0 else ref_cell) for cell in row] for row_index,row in enumerate(PDF_BAD)]
            bad=Table(bad_rows,colWidths=[1.7*reports.inch,4.45*reports.inch],repeatRows=1)
            bad.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),reports._rl(reports.NAVY)),('TEXTCOLOR',(0,0),(-1,0),colors.white),('FONTNAME',(0,0),(-1,0),bold_font),('FONTNAME',(0,1),(-1,-1),regular_font),('FONTSIZE',(0,0),(-1,-1),7.5),('GRID',(0,0),(-1,-1),.35,reports._rl(reports.LINE)),('BACKGROUND',(1,1),(1,1),reports._rl(reports.GREEN_FILL)),('BACKGROUND',(1,2),(1,2),reports._rl(reports.AMBER_FILL)),('BACKGROUND',(1,3),(1,3),reports._rl(reports.RED_FILL))]))
            topo_rows=[[Paragraph(tr(cell),ref_head if row_index==0 else ref_cell) for cell in row] for row_index,row in enumerate(TOPOGRAPHY)]
            topo=Table(topo_rows,colWidths=[1.5*reports.inch,3.95*reports.inch,.7*reports.inch],repeatRows=1)
            topo.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),reports._rl(reports.NAVY)),('TEXTCOLOR',(0,0),(-1,0),colors.white),('FONTNAME',(0,0),(-1,0),bold_font),('FONTNAME',(0,1),(-1,-1),regular_font),('FONTSIZE',(0,0),(-1,-1),7.1),('GRID',(0,0),(-1,-1),.35,reports._rl(reports.LINE)),('ALIGN',(-1,1),(-1,-1),'CENTER')]))
            erss_rows=[[Paragraph(tr(cell),ref_head if row_index==0 else ref_cell) for cell in row] for row_index,row in enumerate(ACTIVE_ERSS)]
            erss=Table(erss_rows,colWidths=[1.7*reports.inch,3.7*reports.inch,.75*reports.inch],repeatRows=1)
            erss.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),reports._rl(reports.NAVY)),('TEXTCOLOR',(0,0),(-1,0),colors.white),('FONTNAME',(0,0),(-1,0),bold_font),('FONTNAME',(0,1),(-1,-1),regular_font),('FONTSIZE',(0,0),(-1,-1),7.1),('GRID',(0,0),(-1,-1),.35,reports._rl(reports.LINE))]))
            topography_block=KeepTogether([
                Paragraph(tr('Randleman topography assessment'),sec),
                topo,
                Paragraph(tr(reports.RANDLEMAN_TOPOGRAPHY_SAFETY),tiny),
                Paragraph(tr(reports.RANDLEMAN_SUPERIOR_NOTE),tiny),
            ])
            active_erss_block=KeepTogether([
                Paragraph(tr('Active CER-AI Randleman / ERSS points'),sec),
                erss,
                Paragraph(tr('Randleman/ERSS is calculated from five independent LASIK inputs. BAD-D and NICE remain separate pathways. Overall ERSS: 0-2 low, 3 moderate, >=4 high; CER-AI does not clear totals >=3.'),tiny),
            ])
            appendix=[Paragraph(tr('CER-AI BAD-D reference points'),sec),bad,Paragraph(tr('BAD-D is read from the Pentacam BAD display and is independent of Randleman/ERSS anterior-topography scoring.'),tiny),topography_block,active_erss_block,Spacer(1,8)]
            story[idx:idx]=appendix
            return orig_build(doc,story,*a,**kw)
        reports.SimpleDocTemplate.build=patched_build
        try:return _orig_pdf(payload)
        finally:reports.SimpleDocTemplate.build=orig_build

def build_docx(payload):
    # Same serialization rule for the temporary helper hooks used by the legacy DOCX builder.
    with _EXPORT_LOCK:
        locale=reports.normalize_locale(payload.get("locale"))
        tr=lambda value:reports.translate_text(value,locale)
        orig_heading=reports._add_heading
        orig_bullet=reports._add_bullet
        state={"suppress_identity":False}
        def patched_bullet(document,text):
            if state["suppress_identity"]:return None
            return orig_bullet(document,text)
        def patched_heading(document,text,level=1):
            if text==tr(IDENTITY_HEADING):
                state["suppress_identity"]=True
                return orig_heading(document,text,level)
            state["suppress_identity"]=False
            if text==tr('Interpretation note'):
                orig_heading(document,tr('CER-AI BAD-D reference points'),1)
                t=document.add_table(rows=1,cols=2);t.style='Table Grid';t.rows[0].cells[0].text=tr('Final BAD-D');t.rows[0].cells[1].text=tr('CER-AI interpretation / action')
                for row in PDF_BAD[1:]:
                    c=t.add_row().cells;c[0].text=tr(row[0]);c[1].text=tr(row[1])
                reports._style_doc_table(t,[1.6,4.25],header=True)
                p=document.add_paragraph(tr('BAD-D is read from the Pentacam BAD display and is independent of Randleman/ERSS anterior-topography scoring.'));p.runs[0].font.size=Pt(8)
                orig_heading(document,tr('Randleman topography assessment'),1)
                topography=document.add_table(rows=1,cols=3);topography.style='Table Grid'
                for j,x in enumerate(TOPOGRAPHY[0]):topography.rows[0].cells[j].text=tr(x)
                for row in TOPOGRAPHY[1:]:
                    c=topography.add_row().cells
                    for j,x in enumerate(row):c[j].text=tr(x)
                reports._style_doc_table(topography,[1.45,3.7,.7],header=True)
                for note_text in (reports.RANDLEMAN_TOPOGRAPHY_SAFETY,reports.RANDLEMAN_SUPERIOR_NOTE):
                    p=document.add_paragraph(tr(note_text));p.runs[0].font.size=Pt(8)
                orig_heading(document,tr('Active CER-AI Randleman / ERSS points'),1)
                e=document.add_table(rows=1,cols=3);e.style='Table Grid'
                for j,x in enumerate(ACTIVE_ERSS[0]):e.rows[0].cells[j].text=tr(x)
                for row in ACTIVE_ERSS[1:]:
                    c=e.add_row().cells
                    for j,x in enumerate(row):c[j].text=tr(x)
                reports._style_doc_table(e,[1.65,3.5,.7],header=True)
                p=document.add_paragraph(tr('Randleman/ERSS is calculated from five independent LASIK inputs. BAD-D and NICE remain separate pathways. Overall ERSS: 0-2 low, 3 moderate, >=4 high; CER-AI does not clear totals >=3.'));p.runs[0].font.size=Pt(8)
            return orig_heading(document,text,level)
        reports._add_heading=patched_heading
        reports._add_bullet=patched_bullet
        try:data=_orig_docx(payload)
        finally:
            reports._add_heading=orig_heading
            reports._add_bullet=orig_bullet
        return _sync_docx_version(data)

reports.build_pdf=build_pdf
reports.build_docx=build_docx
