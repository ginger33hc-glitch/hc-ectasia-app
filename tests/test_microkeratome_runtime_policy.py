import copy
from io import BytesIO

import microkeratome_planning_policy as policy
import pytest
import reports
from docx import Document
from pypdf import PdfReader


def _decision(status="PASS"):
    return {
        "status": status,
        "eyes": [{
            "eye": "OD",
            "status": status,
            "action": "unchanged",
            "hard_stops": [],
            "missing": [],
            "score": {"total": 1},
            "values": {
                "procedure": "LASIK",
                "intended_refractive_pattern": "MYOPIC",
                "pachy_thinnest_um": 540,
                "LASIK_selected_plan": "Plan A",
                "LASIK_flap_um": 100,
                "max_ablation_um": 60,
                "optical_zone_mm": 6.5,
                "transition_zone_mm": 9.0,
            },
            "lasik_selected_plan": "Plan A",
        }],
    }


def _extracted():
    return {"eyes": [{
        "eye": "OD", "K1_D": 42.0, "K1_axis_deg": 110,
        "K2_D": 46.5, "K2_axis_deg": 20, "corneal_diameter_mm": 11.2,
        "table_verified_numeric_fields": [
            "K1_D", "K1_axis_deg", "K2_D", "K2_axis_deg", "corneal_diameter_mm",
        ],
    }]}


def test_runtime_appends_planning_without_changing_ectasia_decision(monkeypatch):
    original = _decision()
    monkeypatch.setattr(policy, "_previous_hc_engine", lambda *args, **kwargs: copy.deepcopy(original))
    out = policy.hc_engine_with_microkeratome_planning(
        _extracted(), 30, {"OD": {"procedure": "LASIK", "flap_um": 100}}, {}, {}
    )
    eye = out["eyes"][0]
    assert out["status"] == original["status"]
    assert eye["status"] == original["eyes"][0]["status"]
    assert eye["action"] == "unchanged"
    assert eye["hard_stops"] == [] and eye["score"] == {"total": 1}
    assert eye["microkeratome_planning"]["primary_hinge"].endswith("110° hinge axis)")
    assert eye["microkeratome_planning"]["status_independent"] is True


def test_runtime_does_not_attach_module_to_non_favorable_case(monkeypatch):
    monkeypatch.setattr(policy, "_previous_hc_engine", lambda *args, **kwargs: _decision("STOP-DEFER"))
    out = policy.hc_engine_with_microkeratome_planning(
        _extracted(), 30, {"OD": {"procedure": "LASIK", "flap_um": 100}}, {}, {}
    )
    assert "microkeratome_planning" not in out["eyes"][0]


@pytest.mark.parametrize("status", ["PASS", "CAUTION"])
def test_pass_and_caution_expose_selected_lasik_and_ml7_plan(monkeypatch, status):
    monkeypatch.setattr(policy, "_previous_hc_engine", lambda *args, **kwargs: _decision(status))
    out = policy.hc_engine_with_microkeratome_planning(
        _extracted(), 30, {"OD": {"procedure": "LASIK", "flap_um": 100}}, {}, {}
    )
    eye = out["eyes"][0]
    values = eye["values"]
    ml7 = eye["microkeratome_planning"]

    assert eye["status"] == status
    assert eye["lasik_selected_plan"] == "Plan A"
    assert (values["LASIK_flap_um"], values["optical_zone_mm"], values["transition_zone_mm"]) == (100, 6.5, 9.0)
    assert ml7["assessment_gate"] == status
    assert ml7["vacuum_ring_mm"] is not None
    assert ml7["vacuum_pressure_mmhg"]
    assert ml7["blade_recommendations"]
    assert ml7["primary_hinge"].endswith("110° hinge axis)")


def test_missing_w2w_or_axis_degrades_planning_only(monkeypatch):
    extracted = _extracted()
    extracted["eyes"][0]["K2_axis_deg"] = None
    extracted["eyes"][0]["corneal_diameter_mm"] = None
    monkeypatch.setattr(policy, "_previous_hc_engine", lambda *args, **kwargs: _decision())
    out = policy.hc_engine_with_microkeratome_planning(
        extracted, 30, {"OD": {"procedure": "LASIK", "flap_um": 100}}, {}, {}
    )
    plan = out["eyes"][0]["microkeratome_planning"]
    assert out["status"] == "PASS"
    assert plan["vacuum_ring_mm"] is None
    assert plan["primary_hinge"] == "Perpendicular to steep axis"
    assert any("horizontal white-to-white (HWTW)" in warning for warning in plan["warnings"])
    assert any("numeric hinge axis cannot be calculated" in warning for warning in plan["warnings"])


def test_unverified_or_non_pentacam_diameter_cannot_drive_ring_selection(monkeypatch):
    extracted = _extracted()
    extracted["eyes"][0]["table_verified_numeric_fields"].remove("corneal_diameter_mm")
    monkeypatch.setattr(policy, "_previous_hc_engine", lambda *args, **kwargs: _decision())
    out = policy.hc_engine_with_microkeratome_planning(
        extracted, 30, {"OD": {"procedure": "LASIK", "flap_um": 100}}, {}, {}
    )
    plan = out["eyes"][0]["microkeratome_planning"]
    assert plan["vacuum_ring_mm"] is None
    assert plan["vacuum_pressure_mmhg"] is None
    assert any("horizontal white-to-white (HWTW)" in warning for warning in plan["warnings"])


def test_planning_is_present_in_pdf_and_word_exports(monkeypatch):
    monkeypatch.setattr(policy, "_previous_hc_engine", lambda *args, **kwargs: _decision())
    decision = policy.hc_engine_with_microkeratome_planning(
        _extracted(), 30, {"OD": {"procedure": "LASIK", "flap_um": 100}}, {}, {}
    )
    payload = {"patient": {"name": "Test Patient"}, "decision": decision, "extracted": _extracted()}
    pdf_text = "\n".join(
        page.extract_text() or "" for page in PdfReader(BytesIO(reports.build_pdf(payload))).pages
    )
    assert "Post-assessment ML7 microkeratome planning" in pdf_text
    assert "Plan A" in pdf_text
    assert "Perpendicular to steep axis" in pdf_text

    document = Document(BytesIO(reports.build_docx(payload)))
    doc_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "Post-assessment ML7 microkeratome planning" in doc_text
    assert "Plan A" in doc_text
    assert "+10 blade; temporal or nasal hinge" in doc_text
