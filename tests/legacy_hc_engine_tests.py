import json
import struct
import unittest
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import Mock, patch

import canonical_engine
import reports
from docx import Document
from docx.shared import Pt
from fastapi.testclient import TestClient
from pypdf import PdfReader

app = canonical_engine.core


def normal_eye(eye="OD", pachy=560, morphology="NORMAL_SYMMETRIC"):
    return {
        "eye": eye,
        "nice_candidates": [{"central_pachy_um": 565, "central_status": "CONFIDENT",
                             "central_landmark": "PUPIL_CENTER_PLUS",
                             "B_Ele_Th_um": 8, "b_ele_th_status": "CONFIDENT",
                             "b_ele_th_landmark": "B_ELE_TH_LABELED_BOX",
                             "b_ele_th_page": "BAD_DISPLAY",
                             "evidence": "Synthetic central 565; BAD Display B. Ele.Th +8"}],
        "screen_types": ["4 Maps", "BAD Display"],
        "keratometry_source": "SHOW_2_EXAMS_TOPOMETRIC_CORNEA_FRONT",
        "quality": "ADEQUATE",
        "pentacam_qs": "OK",
        "missing_or_unreadable": [],
        "table_verified_numeric_fields": list(app.TABLE_NUMERIC_FIELDS),
        "map_fallback_numeric_fields": [],
        "K1_D": 42.0,
        "K2_D": 43.0,
        "Kmax_D": 44.0,
        "pachy_thinnest_um": pachy,
        "BAD_D": 1.0,
        "Df": 1.0,
        "Db": 1.0,
        "Dp": 1.0,
        "Dt": -0.5,
        "Da": 0.2,
        "PPI_avg": 0.9,
        "PPI_min": 0.7,
        "PPI_max": 1.1,
        "ARTmax_um": 500,
        "ISV": 20,
        "IVA": 0.1,
        "KI": 1.0,
        "CKI": 1.0,
        "IHD": 0.01,
        "I_S": 0.5,
        "KISA": 10,
        "IHA": 2.0,
        "Rmin_mm": 7.6,
        "anterior_elevation_thinnest_um": 2.0,
        "posterior_elevation_thinnest_um": 4.0,
        "thinnest_x_mm": 0.2,
        "thinnest_y_mm": -0.3,
        "corneal_volume_mm3": 60.0,
        "RMS_HOA_um": 0.2,
        "vertical_coma_um": 0.05,
        "Kmean_D": 42.5,
        "total_RMS_um": 0.3,
        "spherical_aberration_um": 0.1,
        "morphology": morphology,
        "morphology_evidence": ["Visible symmetric bowtie"],
        "asymmetric_bow_tie": "NO",
        "srax": "NO",
        "srax_deg": 0,
        "inferior_opposite_steepening_D": 0,
        "anterior_pattern": "REASSURING",
        "posterior_pattern": "REASSURING",
        "anterior_curvature_map_visible": "YES",
        "anterior_curvature_map_type": "AXIAL_SAGITTAL_FRONT",
        "anterior_curvature_map_location": "UPPER_LEFT",
    }


def plan(procedure="PRK", sphere=-3.0, cylinder=0.0, ablation=60, flap=None):
    return {
        "prior": "no",
        "procedure": procedure,
        "manifest_sphere_D": sphere,
        "manifest_cylinder_magnitude_D": cylinder,
        "intended_sphere_D": sphere,
        "intended_cylinder_magnitude_D": cylinder,
        "ablation_um": ablation,
        "flap_um": flap,
        "laser_platform": "Alcon EX500",
        "optical_zone_mm": 6.0,
        "transition_zone_mm": None,
        "transition_zone_not_applicable": "yes",
        "stable": "yes",
        "progression": "no",
        "cdva_below_20_20": "no",
        "enhancement_anticipated": "no",
    }


def card_correction(
    eye="OD", sphere=-4.5, cylinder=-3.5, axis=170,
    sphere_cylinder_status="CONFIDENT", axis_status="CONFIDENT",
):
    return {
        "eye": eye,
        "source_document": "EXCIMER_LASER_FOLLOW_UP_CARD",
        "source_label": "DUZELTME_MIKTARI",
        "sphere_D": sphere,
        "cylinder_D": cylinder,
        "axis_deg": axis,
        "sphere_cylinder_status": sphere_cylinder_status,
        "axis_status": axis_status,
        "raw_text": f"{sphere} ({cylinder} x {axis})",
        "missing_or_unreadable": [],
    }


MODIFIERS = {
    "eye_rubbing": "no",
    "family_history": "no",
    "inter_eye_asymmetry": "no",
    "pregnancy_nursing": "no",
    "collagen_tissue_disease": "no",
    "drug_usage": "no",
    "dry_eye": "no",
    "systemic_disease": "no",
    "contact_lens_type": "NONE",
    "contact_lens_discontinuation_days": None,
    "assessed_eyes": ["OD"],
}


def document_context(
    patient_id="CER-AI-1", patient_age_years=26, exam_date="2026-08-25", qs="OK",
    patient_name="Test Patient", laterality="BOTH",
):
    name_parts = str(patient_name or "").split(maxsplit=1)
    first_name = name_parts[0] if name_parts else None
    last_name = name_parts[1] if len(name_parts) > 1 else None
    return {
        "document_type": "PENTACAM_TOPOGRAPHY", "patient_id": patient_id,
        "patient_first_name": first_name, "patient_last_name": last_name,
        "patient_name": patient_name,
        "patient_name_source": "PENTACAM_FIRST_LAST_NAME_FIELDS" if patient_name else "UNREADABLE",
        "patient_age_years": patient_age_years, "exam_date": exam_date,
        "exam_time": "10:00", "laterality": laterality, "pentacam_qs": qs,
        "missing_or_unreadable": [], "source_filename": "pentacam.png",
    }


class TestSafetyGates(unittest.TestCase):
    def test_prior_surgery_short_circuits_virgin_engine_even_if_thin(self):
        prior_plan = plan()
        prior_plan["prior"] = "yes"
        result = app.assess_eye(normal_eye(pachy=430), prior_plan, 35, MODIFIERS)
        self.assertEqual(result["status"], "POST-REFRACTIVE PATHWAY REQUIRED")
        self.assertEqual(result["score"]["rows"], {})
        self.assertEqual(result["hard_stops"], [])

    def test_hc_sphere_hard_stops_and_exact_boundaries(self):
        minus_11 = app.assess_eye(normal_eye(), plan(sphere=-11), 35, MODIFIERS)
        minus_10 = app.assess_eye(normal_eye(), plan(sphere=-10), 35, MODIFIERS)
        plus_7 = app.assess_eye(normal_eye(), plan(sphere=7), 35, MODIFIERS)
        plus_6 = app.assess_eye(normal_eye(), plan(sphere=6), 35, MODIFIERS)
        self.assertTrue(any("<−10.00" in item for item in minus_11["hard_stops"]))
        self.assertFalse(any("<−10.00" in item for item in minus_10["hard_stops"]))
        self.assertTrue(any(">+6.00" in item for item in plus_7["hard_stops"]))
        self.assertFalse(any(">+6.00" in item for item in plus_6["hard_stops"]))

    def test_final_keratometry_uses_intended_mrse_and_exact_boundaries_are_allowed(self):
        formula_eye = normal_eye()
        formula_eye["Kmean_D"] = 42.5
        formula = app.assess_eye(
            formula_eye, plan(sphere=-4.5, cylinder=3.5), 35, MODIFIERS
        )
        self.assertEqual(formula["values"]["intended_MRSE_D"], -6.25)
        self.assertEqual(formula["values"]["estimated_final_Kmean_D"], 37.5)

        lower_eye = normal_eye()
        lower_eye["Kmean_D"] = 44.0
        lower = app.assess_eye(lower_eye, plan(sphere=-10.0, cylinder=0.0), 35, MODIFIERS)
        upper_eye = normal_eye()
        upper_eye["Kmean_D"] = 43.2
        upper = app.assess_eye(upper_eye, plan(sphere=6.0, cylinder=0.0), 35, MODIFIERS)
        self.assertEqual(lower["values"]["estimated_final_Kmean_D"], 36.0)
        self.assertEqual(upper["values"]["estimated_final_Kmean_D"], 48.0)
        self.assertFalse(any("final-keratometry" in item for item in lower["hard_stops"]))
        self.assertFalse(any("final-keratometry" in item for item in upper["hard_stops"]))

    def test_final_keratometry_outside_36_to_48_is_independent_hard_stop(self):
        flat_eye = normal_eye()
        flat_eye["Kmean_D"] = 43.5
        flat = app.assess_eye(flat_eye, plan(sphere=-9.5, cylinder=0.0), 35, MODIFIERS)
        steep_eye = normal_eye()
        steep_eye["Kmean_D"] = 43.3
        steep = app.assess_eye(steep_eye, plan(sphere=6.0, cylinder=0.0), 35, MODIFIERS)
        self.assertAlmostEqual(flat["values"]["estimated_final_Kmean_D"], 35.9)
        self.assertAlmostEqual(steep["values"]["estimated_final_Kmean_D"], 48.1)
        self.assertEqual(flat["status"], "STOP-DEFER")
        self.assertEqual(steep["status"], "STOP-DEFER")
        self.assertTrue(any("<36.00" in item for item in flat["hard_stops"]))
        self.assertTrue(any(">48.00" in item for item in steep["hard_stops"]))

    def test_missing_kmean_prohibits_pass_instead_of_inferring_from_k1_k2(self):
        eye = normal_eye()
        eye["Kmean_D"] = None
        result = app.assess_eye(eye, plan(), 35, MODIFIERS)
        self.assertIsNone(result["values"]["estimated_final_Kmean_D"])
        self.assertIn("preoperative Kmean for final keratometry safety calculation", result["missing"])
        self.assertNotEqual(result["status"], "PASS")

    def test_manifest_refraction_not_intended_plan_drives_lasik_mrse(self):
        p = plan("LASIK", sphere=-9, cylinder=2, ablation=70, flap=100)
        p["manifest_sphere_D"] = -1
        p["manifest_cylinder_magnitude_D"] = 0
        result = app.assess_eye(normal_eye(), p, 35, MODIFIERS)
        self.assertEqual(result["values"]["MRSE_D"], -1)
        self.assertEqual(result["score"]["rows"]["MRSE"], 0)

    def test_negative_ablation_is_rejected_and_not_used(self):
        result = app.assess_eye(normal_eye(), plan(ablation=-50), 35, MODIFIERS)
        self.assertIsNone(result["values"]["max_ablation_um"])
        self.assertIsNone(result["values"]["PRK_RST_um"])
        self.assertNotEqual(result["status"], "PASS")

    def test_i_s_merge_does_not_create_definite_disease_override(self):
        eye = normal_eye()
        eye["I_S"] = 1.4
        merged = app.merge_extractions([{"eyes": [eye], "treatment_corrections": [], "global_warnings": []}])
        self.assertEqual(merged["eyes"][0]["morphology"], "NORMAL_SYMMETRIC")
        result = app.assess_eye(merged["eyes"][0], plan("LASIK", ablation=60, flap=100), 35, MODIFIERS)
        self.assertEqual(result["topography_classification"]["scoring_category"], "ABNORMAL_ECTATIC")
        self.assertFalse(any("Definite KC" in item for item in result["hard_stops"]))

    def test_single_eye_never_yields_overall_pass(self):
        extracted = {"eyes": [normal_eye("OD")], "global_warnings": []}
        result = app.hc_engine(extracted, 35, {"OD": plan()}, MODIFIERS)
        self.assertEqual(result["status"], "DATA INSUFFICIENT")
        self.assertTrue(any("Both OD and OS" in item for item in result["critical_input_issues"]))

    def test_identity_and_date_conflicts_block_but_non_ok_qs_warns(self):
        first = {"document_context": document_context("A", qs="OK"), "eyes": [normal_eye("OD")], "treatment_corrections": [], "global_warnings": []}
        second_context = document_context("B", exam_date="2026-08-26", qs="NOT_OK", patient_name="Different Patient")
        second_context["source_filename"] = "other.png"
        second = {"document_context": second_context, "eyes": [normal_eye("OS")], "treatment_corrections": [], "global_warnings": []}
        merged = app.merge_extractions([first, second])
        self.assertTrue(any("conflicting patient IDs" in item for item in merged["identity_warnings"]))
        self.assertTrue(any("different patient names" in item for item in merged["identity_warnings"]))
        self.assertTrue(any("Conflicting Pentacam" in item for item in merged["critical_input_issues"]))
        self.assertFalse(any("QS" in item for item in merged["critical_input_issues"]))
        decision = app.hc_engine(merged, 35, {"OD": plan(), "OS": plan()}, MODIFIERS)
        self.assertTrue(any("QS is NOT_OK" in item for item in decision["source_quality_warnings"]))

    def test_different_extracted_ids_do_not_block_when_name_and_age_match(self):
        od_context = document_context(
            patient_id="P-77", patient_name="Same Patient", patient_age_years=35, laterality="OD"
        )
        od_context["source_filename"] = "od.png"
        os_context = document_context(
            patient_id="SCAN-88", patient_name="Same Patient", patient_age_years=35, laterality="OS"
        )
        os_context["source_filename"] = "os.png"
        merged = app.merge_extractions([
            {"document_context": od_context, "eyes": [normal_eye("OD")], "treatment_corrections": [], "global_warnings": []},
            {"document_context": os_context, "eyes": [normal_eye("OS")], "treatment_corrections": [], "global_warnings": []},
        ])
        self.assertFalse(any("Conflicting patient IDs" in item for item in merged["critical_input_issues"]))
        self.assertTrue(any("different patient-ID strings" in item for item in merged["identity_warnings"]))
        self.assertEqual(merged["derived_age_years"], 35)

    def test_patient_id_prompt_rejects_scan_and_measurement_numbers(self):
        self.assertIn("Never use an examination number, measurement number, scan number", app.PROMPT)
        self.assertIn("return patient_id=null", app.PROMPT)

    def test_matching_od_os_pentacam_identity_and_printed_age_are_used(self):
        od_context = document_context(patient_id="P-77", patient_age_years=35, laterality="OD")
        od_context["source_filename"] = "od.png"
        os_context = document_context(patient_id="P-77", patient_age_years=35, laterality="OS")
        os_context["source_filename"] = "os.png"
        merged = app.merge_extractions([
            {"document_context": od_context, "eyes": [normal_eye("OD")], "treatment_corrections": [], "global_warnings": []},
            {"document_context": os_context, "eyes": [normal_eye("OS")], "treatment_corrections": [], "global_warnings": []},
        ])
        self.assertEqual(merged["derived_age_years"], 35)
        self.assertFalse(any("same patient" in item for item in merged["critical_input_issues"]))
        decision = app.hc_engine(
            merged, None, {"OD": plan(), "OS": plan()}, MODIFIERS, {"id": "P-77"}
        )
        self.assertEqual([eye["values"]["age_years"] for eye in decision["eyes"]], [35, 35])

    def test_conflicting_pentacam_ages_require_one_manual_patient_age(self):
        od_context = document_context(patient_age_years=35, laterality="OD")
        os_context = document_context(patient_age_years=36, laterality="OS")
        merged = app.merge_extractions([
            {"document_context": od_context, "eyes": [normal_eye("OD")], "treatment_corrections": [], "global_warnings": []},
            {"document_context": os_context, "eyes": [normal_eye("OS")], "treatment_corrections": [], "global_warnings": []},
        ])
        self.assertNotIn("derived_age_years", merged)
        self.assertEqual(merged["patient_age_conflict_values"], [35, 36])
        self.assertFalse(any("patient ages" in item.lower() for item in merged["critical_input_issues"]))
        without_confirmation = app.hc_engine(
            merged, None, {"OD": plan(), "OS": plan()}, MODIFIERS
        )
        self.assertTrue(all("age" in eye["missing"] for eye in without_confirmation["eyes"]))
        confirmed = app.hc_engine(
            merged, 35, {"OD": plan(), "OS": plan()}, MODIFIERS
        )
        self.assertFalse(any("age" in eye["missing"] for eye in confirmed["eyes"]))

    def test_od_os_sources_without_shared_readable_identifier_cannot_be_verified(self):
        od_context = document_context(patient_id="P-77", patient_name=None, laterality="OD")
        os_context = document_context(patient_id=None, patient_name="Test Patient", laterality="OS")
        merged = app.merge_extractions([
            {"document_context": od_context, "eyes": [normal_eye("OD")], "treatment_corrections": [], "global_warnings": []},
            {"document_context": os_context, "eyes": [normal_eye("OS")], "treatment_corrections": [], "global_warnings": []},
        ])
        self.assertTrue(any("could not be confirmed as the same patient" in item for item in merged["identity_warnings"]))

    def test_pentacam_name_uses_only_first_and_last_name_fields(self):
        context = document_context(patient_name="Wrong Header Name")
        context["patient_first_name"] = "Hasan"
        context["patient_last_name"] = "Can"
        normalized = app.normalize_document_context_identity(context)
        self.assertEqual(normalized["patient_name"], "Hasan Can")
        self.assertEqual(normalized["patient_name_source"], "PENTACAM_FIRST_LAST_NAME_FIELDS")
        self.assertIn('explicitly labeled "Last Name" and "First Name"', app.PROMPT)
        self.assertIn("Never use a physician", app.PROMPT)

    def test_unreadable_name_warns_but_eye_analyses_are_not_suppressed(self):
        od_context = document_context(patient_id=None, patient_name=None, laterality="OD")
        os_context = document_context(patient_id=None, patient_name=None, laterality="OS")
        od_context["source_filename"] = "od.png"
        os_context["source_filename"] = "os.png"
        merged = app.merge_extractions([
            {"document_context": od_context, "eyes": [normal_eye("OD")], "treatment_corrections": [], "global_warnings": []},
            {"document_context": os_context, "eyes": [normal_eye("OS")], "treatment_corrections": [], "global_warnings": []},
        ])
        decision = app.hc_engine(merged, None, {"OD": plan(), "OS": plan()}, MODIFIERS)
        self.assertEqual(len(decision["eyes"]), 2)
        self.assertNotEqual(decision["status"], "DATA INSUFFICIENT")
        self.assertEqual(decision["identity_verification"], "NOT VERIFIED — SURGEON CONFIRMATION REQUIRED")
        self.assertTrue(decision["identity_warnings"])
        self.assertFalse(any("Patient identity" in item for item in decision["critical_input_issues"]))

    def test_entered_name_may_follow_first_last_or_last_first_order(self):
        contexts = [document_context(laterality="OD"), document_context(laterality="OS")]
        contexts[0]["source_filename"] = "od.png"
        contexts[1]["source_filename"] = "os.png"
        merged = app.merge_extractions([
            {"document_context": contexts[0], "eyes": [normal_eye("OD")], "treatment_corrections": [], "global_warnings": []},
            {"document_context": contexts[1], "eyes": [normal_eye("OS")], "treatment_corrections": [], "global_warnings": []},
        ])
        decision = app.hc_engine(
            merged, None, {"OD": plan(), "OS": plan()}, MODIFIERS, {"name": "Patient Test"}
        )
        self.assertFalse(any("entered patient name" in item for item in decision["identity_warnings"]))

    def test_date_of_birth_is_removed_and_manual_age_is_requested_only_after_unreadable_age(self):
        html = Path("static/index.html").read_text(encoding="utf-8")
        self.assertNotIn('id="date_of_birth"', html)
        self.assertIn('id="manualAgeRow" hidden', html)
        self.assertIn("Patient age could not be read from the Pentacam image", html)
        context_schema = app.SCHEMA["properties"]["document_context"]["properties"]
        self.assertNotIn("date_of_birth", context_schema)
        self.assertIn("patient_age_years", context_schema)

    def test_clinical_eligibility_modifier_blocks_pass_without_score_points(self):
        modifiers = dict(MODIFIERS, pregnancy_nursing="yes")
        result = app.assess_eye(normal_eye(), plan(), 35, modifiers)
        self.assertEqual(result["status"], "STOP-DEFER")
        self.assertEqual(result["score"]["total"], 0)


class TestBoundaries(unittest.TestCase):
    def test_bad_display_boundaries(self):
        self.assertEqual(app.bad_classification(1.6, final=True), "NORMAL")
        self.assertEqual(app.bad_classification(1.6001, final=True), "SUSPICIOUS")
        self.assertEqual(app.bad_classification(2.5999, final=True), "SUSPICIOUS")
        self.assertEqual(app.bad_classification(2.6, final=True), "ABNORMAL")
        self.assertEqual(app.bad_classification(1.6), "SUSPICIOUS")
        self.assertEqual(app.bad_classification(2.6), "ABNORMAL")

    def test_prk_cct_480_not_hard_stop_but_scores_caution(self):
        result = app.assess_eye(normal_eye(pachy=480), plan(ablation=120), 35, MODIFIERS)
        self.assertEqual(result["values"]["PRK_RST_um"], 310)
        self.assertEqual(result["score"]["total"], 3)
        self.assertEqual(result["status"], "STOP-DEFER")
        self.assertEqual(result["hard_stops"], [])

    def test_prk_cct_479_is_hard_stop_even_with_other_missing_data(self):
        eye = normal_eye(pachy=479)
        eye["BAD_D"] = None
        result = app.assess_eye(eye, plan(), 35, MODIFIERS)
        self.assertEqual(result["status"], "STOP-DEFER")
        self.assertTrue(any("<480" in item for item in result["hard_stops"]))
        self.assertIn("BAD_D", result["missing"])

    def test_prk_rst_310_allowed_by_structural_rule(self):
        result = app.assess_eye(normal_eye(pachy=520), plan(ablation=160), 35, MODIFIERS)
        self.assertEqual(result["values"]["PRK_epithelium_um"], 50)
        self.assertEqual(result["values"]["PRK_RST_um"], 310)
        self.assertFalse(any("RST" in item for item in result["hard_stops"]))

    def test_prk_rst_below_310_is_hard_stop(self):
        result = app.assess_eye(normal_eye(pachy=520), plan(ablation=161), 35, MODIFIERS)
        self.assertEqual(result["status"], "STOP-DEFER")
        self.assertTrue(any("RST <310" in item for item in result["hard_stops"]))

    def test_lasik_rsb_failure_triggers_automatic_safe_fallback(self):
        allowed = app.assess_eye(
            normal_eye(pachy=520), plan("LASIK", ablation=120, flap=100), 35, MODIFIERS
        )
        stopped = app.assess_eye(
            normal_eye(pachy=520), plan("LASIK", ablation=121, flap=100), 35, MODIFIERS
        )
        self.assertEqual(allowed["lasik_planning_sequence"][0]["LASIK_RSB_um"], 300)
        self.assertEqual(stopped["lasik_planning_sequence"][0]["LASIK_RSB_um"], 299)
        self.assertEqual(stopped["lasik_planning_sequence"][0]["status"], "STOP-DEFER")
        self.assertEqual(allowed["lasik_selected_plan"], "Plan B")
        self.assertEqual(stopped["lasik_selected_plan"], "Plan B")
        self.assertGreaterEqual(allowed["values"]["LASIK_RSB_um"], 300)
        self.assertGreaterEqual(stopped["values"]["LASIK_RSB_um"], 300)

    def test_lasik_exact_510_boundary_receives_confirmed_score(self):
        result = app.assess_eye(
            normal_eye(pachy=510), plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["score"]["rows"]["pachymetry"], 0)
        self.assertNotIn("pachymetry boundary at exactly 510 um", result["missing"])

    def test_i_s_1_4_uses_published_erss_abnormal_pattern_without_disease_override(self):
        eye = normal_eye()
        eye["I_S"] = 1.4
        result = app.assess_eye(
            eye, plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["topography_classification"]["scoring_category"], "ABNORMAL_ECTATIC")
        self.assertEqual(result["score"]["rows"]["topography"], 4)
        self.assertEqual(result["status"], "STOP-DEFER")
        self.assertFalse(any("Definite KC" in item for item in result["hard_stops"]))

    def test_srax_20_degrees_uses_published_erss_sra_category(self):
        eye = normal_eye()
        eye["srax"] = "UNCERTAIN"
        eye["srax_deg"] = 20
        result = app.assess_eye(
            eye, plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["topography_classification"]["scoring_category"], "INFERIOR_STEEPENING_SRA")
        self.assertEqual(result["score"]["rows"]["topography"], 3)
        self.assertEqual(result["score"]["total"], 3)
        self.assertEqual(result["status"], "CAUTION")

    def test_minimal_axis_deviation_is_not_scored_as_srax(self):
        eye = normal_eye()
        eye["morphology"] = "INFERIOR_STEEPENING_SRA"
        eye["srax"] = "YES"
        eye["srax_deg"] = 5
        result = app.assess_eye(
            eye, plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["topography_classification"]["scoring_category"], "UNCERTAIN")
        self.assertIsNone(result["score"]["rows"]["topography"])
        self.assertNotEqual(result["status"], "PASS")

    def test_srax_below_20_degrees_is_not_scored(self):
        eye = normal_eye()
        eye["morphology"] = "INFERIOR_STEEPENING_SRA"
        eye["srax"] = "YES"
        eye["srax_deg"] = 19.9
        result = app.assess_eye(
            eye, plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["topography_classification"]["scoring_category"], "UNCERTAIN")
        self.assertNotEqual(result["status"], "PASS")

    def test_unquantified_visual_srax_label_is_not_scored(self):
        eye = normal_eye()
        eye["morphology"] = "INFERIOR_STEEPENING_SRA"
        eye["srax"] = "YES"
        eye["srax_deg"] = None
        result = app.assess_eye(
            eye, plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["topography_classification"]["scoring_category"], "UNCERTAIN")
        self.assertNotEqual(result["status"], "PASS")

    def test_quantified_inferior_steepening_alternative_uses_published_category(self):
        eye = normal_eye()
        eye["morphology"] = "UNCERTAIN"
        eye["I_S"] = 1.3
        eye["inferior_opposite_steepening_D"] = 1.0
        result = app.assess_eye(
            eye, plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["topography_classification"]["scoring_category"], "INFERIOR_STEEPENING_SRA")
        self.assertEqual(result["score"]["rows"]["topography"], 3)

    def test_definite_ectatic_morphology_is_not_downgraded_by_srax_fields(self):
        eye = normal_eye(morphology="ABNORMAL_ECTATIC")
        eye["anterior_curvature_map_visible"] = "YES"
        eye["morphology_confidence"] = "HIGH"
        eye["erss_source_read"] = "DEDICATED_CURVATURE_PASS"
        eye["srax"] = "YES"
        eye["srax_deg"] = 5
        result = app.assess_eye(
            eye, plan("LASIK", sphere=-3, ablation=100, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["topography_classification"]["scoring_category"], "ABNORMAL_ECTATIC")
        self.assertEqual(result["status"], "STOP-DEFER")

    def test_ablation_estimate_uses_zone_specific_ex500_conventions(self):
        zone_6 = plan(ablation=None)
        zone_65 = dict(zone_6, optical_zone_mm=6.5)
        zone_7 = dict(zone_6, optical_zone_mm=7.0)
        other_platform = dict(zone_65, laser_platform="Other platform")

        zone_6_result = app.assess_eye(normal_eye(), zone_6, 35, MODIFIERS)
        zone_65_result = app.assess_eye(normal_eye(), zone_65, 35, MODIFIERS)
        zone_7_result = app.assess_eye(normal_eye(), zone_7, 35, MODIFIERS)
        other_platform_result = app.assess_eye(normal_eye(), other_platform, 35, MODIFIERS)

        self.assertEqual(zone_6_result["values"]["max_ablation_um"], 36.0)
        self.assertEqual(zone_65_result["values"]["max_ablation_um"], 45.0)
        self.assertAlmostEqual(zone_7_result["values"]["max_ablation_um"], 48.99)
        self.assertIn("15 µm/D", zone_65_result["warnings"][0])
        self.assertIn("16.33 µm/D", zone_7_result["warnings"][0])
        self.assertIsNone(other_platform_result["values"]["max_ablation_um"])
        self.assertEqual(other_platform_result["status"], "DATA INSUFFICIENT")

    def test_hyperopic_plan_requires_actual_ablation_and_review(self):
        estimated = app.assess_eye(normal_eye(), plan(sphere=2.0, cylinder=0.0, ablation=None), 35, MODIFIERS)
        actual = app.assess_eye(normal_eye(), plan(sphere=2.0, cylinder=0.0, ablation=40), 35, MODIFIERS)
        self.assertIsNone(estimated["values"]["max_ablation_um"])
        self.assertNotEqual(estimated["status"], "PASS")
        self.assertEqual(actual["status"], "CAUTION")
        self.assertFalse(any("treatment cutoff" in item for item in actual["hard_stops"]))

    def test_hyperopic_report_is_populated_and_contains_case_specific_surgeon_attention(self):
        result = app.assess_eye(
            normal_eye(), plan("LASIK", sphere=3.0, cylinder=1.0, ablation=55, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["values"]["intended_refractive_pattern"], "HYPEROPIC")
        self.assertEqual(result["values"]["intended_principal_meridians_D"], [3.0, 2.0])
        self.assertIsNotNone(result["values"]["LASIK_RSB_um"])
        self.assertTrue(any("cycloplegic" in item for item in result["surgeon_attention"]))
        self.assertTrue(any("PMD" in item for item in result["surgeon_attention"]))
        self.assertTrue(any("+6.00 D sphere" in item for item in result["surgeon_attention"]))

    def test_mixed_plan_does_not_use_mrse_or_kmean_as_surgical_load_reassurance(self):
        result = app.assess_eye(
            normal_eye(), plan("LASIK", sphere=3.0, cylinder=6.0, ablation=90, flap=100), 35, MODIFIERS
        )
        self.assertEqual(result["values"]["intended_refractive_pattern"], "MIXED_ASTIGMATISM")
        self.assertEqual(result["values"]["intended_principal_meridians_D"], [3.0, -3.0])
        self.assertEqual(result["values"]["intended_MRSE_D"], 0.0)
        self.assertIsNone(result["values"]["estimated_final_Kmean_D"])
        self.assertIsNotNone(result["values"]["LASIK_RSB_um"])
        self.assertEqual(result["status"], "DATA INSUFFICIENT")
        self.assertTrue(any("near-zero MRSE" in item for item in result["surgeon_attention"]))
        self.assertTrue(any("meridional K" in item for item in result["surgeon_attention"]))

    def test_mixed_classification_is_invariant_to_plus_cylinder_entry(self):
        signed = plan("LASIK", sphere=None, cylinder=None, ablation=90, flap=100)
        signed.update({
            "manifest_entered_sphere_D": -3.0, "manifest_cylinder_signed_D": 6.0,
            "intended_entered_sphere_D": -3.0, "intended_cylinder_signed_D": 6.0,
            "entered_axis_deg": 90,
        })
        result = app.assess_eye(normal_eye(), signed, 35, MODIFIERS)
        self.assertEqual(result["values"]["intended_sphere_D"], 3.0)
        self.assertEqual(result["values"]["intended_cylinder_magnitude_D"], 6.0)
        self.assertEqual(result["values"]["intended_refractive_pattern"], "MIXED_ASTIGMATISM")

    def test_mixed_lasik_label_limits_are_reported_as_review_not_ectasia_score(self):
        young = app.assess_eye(
            normal_eye(), plan("LASIK", sphere=3.0, cylinder=6.0, ablation=90, flap=100), 20, MODIFIERS
        )
        outside_cylinder = app.assess_eye(
            normal_eye(), plan("LASIK", sphere=3.0, cylinder=6.25, ablation=90, flap=100), 35, MODIFIERS
        )
        self.assertTrue(any("labeled age range" in item for item in young["reasons"]))
        self.assertTrue(any("labeled range" in item for item in outside_cylinder["reasons"]))
        self.assertFalse(any("mixed" in item.lower() for item in outside_cylinder["hard_stops"]))

    def test_hyperopic_mixed_attention_is_routed_to_browser_pdf_and_word_reports(self):
        result = app.assess_eye(
            normal_eye(), plan("LASIK", sphere=3.0, cylinder=6.0, ablation=90, flap=100), 35, MODIFIERS
        )
        findings = dict(reports._findings(result))
        self.assertIn("Surgeon attention - hyperopic/mixed pathway", findings)
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn("Surgeon attention — hyperopic/mixed pathway", html)

    def test_prk_mitomycin_c_guidance_uses_myopia_magnitude_boundary(self):
        required_boundary = app.assess_eye(
            normal_eye(), plan("PRK", sphere=-4.0, cylinder=0.0, ablation=60), 35, MODIFIERS
        )
        required_greater_myopia = app.assess_eye(
            normal_eye(), plan("PRK", sphere=-5.0, cylinder=0.0, ablation=60), 35, MODIFIERS
        )
        recommended_lower_myopia = app.assess_eye(
            normal_eye(), plan("PRK", sphere=-3.99, cylinder=0.0, ablation=60), 35, MODIFIERS
        )

        self.assertIn("REQUIRED", required_boundary["prk_mitomycin_c_guidance"][0])
        self.assertIn("-4.00 D or -5.00 D", required_boundary["prk_mitomycin_c_guidance"][0])
        self.assertIn("REQUIRED", required_greater_myopia["prk_mitomycin_c_guidance"][0])
        self.assertIn("RECOMMENDED", recommended_lower_myopia["prk_mitomycin_c_guidance"][0])

    def test_prk_mitomycin_c_guidance_requires_hyperopic_use_and_does_not_affect_lasik(self):
        hyperopic_prk = app.assess_eye(
            normal_eye(), plan("PRK", sphere=2.0, cylinder=0.0, ablation=40), 35, MODIFIERS
        )
        myopic_lasik = app.assess_eye(
            normal_eye(), plan("LASIK", sphere=-5.0, cylinder=0.0, ablation=70, flap=100), 35, MODIFIERS
        )

        self.assertEqual(
            hyperopic_prk["prk_mitomycin_c_guidance"],
            ["Mitomycin-C use is REQUIRED for hyperopic PRK."],
        )
        self.assertEqual(myopic_lasik["prk_mitomycin_c_guidance"], [])

    def test_prk_mitomycin_c_guidance_is_routed_to_all_report_formats(self):
        result = app.assess_eye(
            normal_eye(), plan("PRK", sphere=-4.0, cylinder=0.0, ablation=60), 35, MODIFIERS
        )
        findings = dict(reports._findings(result))
        self.assertIn("PRK Mitomycin-C guidance", findings)
        self.assertIn("REQUIRED", findings["PRK Mitomycin-C guidance"][0])

        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('listBlock("PRK Mitomycin-C guidance",r.prk_mitomycin_c_guidance)', html)


class TestScoringAndCompleteness(unittest.TestCase):
    def test_extraction_contract_prioritizes_labeled_pentacam_numeric_fields(self):
        eye_schema = app.SCHEMA["properties"]["eyes"]["items"]
        self.assertIn("table_verified_numeric_fields", eye_schema["required"])
        self.assertIn("map_fallback_numeric_fields", eye_schema["required"])
        self.assertIn("keratometry_source", eye_schema["required"])
        self.assertIn("PENTACAM NUMERIC-SOURCE RULE", app.PROMPT)
        self.assertIn('screen whose visible title is "Show 2 Exams Topometric"', app.PROMPT)
        self.assertIn('upper parameter panel explicitly headed "Cornea Front"', app.PROMPT)
        self.assertIn("Kmean_D is the value printed in the Km row", app.PROMPT)
        self.assertIn("Never substitute a generic map spot value", app.PROMPT)
        self.assertIn("Only the categorical fields", app.PROMPT)

    def test_reassuring_prk_can_pass(self):
        result = app.assess_eye(normal_eye(), plan(), 35, MODIFIERS)
        self.assertEqual(result["score"]["total"], 0)
        self.assertEqual(result["status"], "PASS")

    def test_lasik_score_two_has_no_score_escalation(self):
        eye = normal_eye(morphology="ASYMMETRIC_BOWTIE")
        eye["asymmetric_bow_tie"] = "YES"
        eye["inferior_opposite_steepening_D"] = 0.7
        eye["Kmean_D"] = 44.0
        treatment = plan("LASIK", sphere=-8.5, cylinder=1.0, ablation=100, flap=100)
        treatment["correction_axis_deg"] = 180
        treatment["manifest_normalized_axis_deg"] = 180
        treatment["intended_normalized_axis_deg"] = 180
        result = app.assess_eye(
            eye, treatment, 28, MODIFIERS
        )
        self.assertEqual(result["score"]["total"], 2)
        self.assertEqual(result["status"], "PASS")

    def test_missing_critical_tomography_prohibits_pass(self):
        eye = normal_eye()
        eye["ARTmax_um"] = None
        result = app.assess_eye(eye, plan(), 35, MODIFIERS)
        self.assertEqual(result["status"], "DATA INSUFFICIENT")
        self.assertIn("ARTmax_um", result["missing"])

    def test_prk_pta_outside_cohort_envelope_is_labeled_not_recast_as_validated_cutoff(self):
        result = app.assess_eye(normal_eye(pachy=560), plan(ablation=150), 35, MODIFIERS)
        self.assertGreater(result["values"]["PRK_PTA_percent"], 35.28)
        self.assertTrue(any("evidence-gap" in item for item in result["surgical_load_flags"]))
        self.assertEqual(result["status"], "CAUTION")

    def test_concordant_cross_sectional_tomography_flags_prohibit_pass(self):
        eye = normal_eye(pachy=530)
        eye.update(ARTmax_um=400, Dt=-0.10, Da=0.70)
        result = app.assess_eye(eye, plan(), 35, MODIFIERS)
        self.assertEqual(result["tomography_review"]["status"], "CONCERN FLAGS")
        self.assertEqual(len(result["tomography_review"]["cross_sectional_flags"]), 4)
        self.assertEqual(result["status"], "CAUTION")

    def test_limited_image_quality_allows_pass_with_warning(self):
        eye = normal_eye()
        eye["quality"] = "LIMITED"
        result = app.assess_eye(eye, plan(), 35, MODIFIERS)
        self.assertEqual(result["status"], "PASS")
        self.assertNotIn("adequate-quality tomography/topography", result["missing"])
        extracted = {"eyes": [eye, normal_eye("OS")], "global_warnings": []}
        decision = app.hc_engine(extracted, 35, {"OD": plan(), "OS": plan()}, MODIFIERS)
        self.assertTrue(any("source image quality is LIMITED" in item for item in decision["source_quality_warnings"]))

        eye_without_bad_d = dict(eye, BAD_D=None)
        incomplete = app.assess_eye(eye_without_bad_d, plan(), 35, MODIFIERS)
        self.assertEqual(incomplete["status"], "DATA INSUFFICIENT")
        self.assertIn("BAD_D", incomplete["missing"])

    def test_limited_ancillary_page_does_not_poison_adequate_same_eye_source(self):
        adequate = normal_eye()
        adequate["_source_filename"] = "decision.png"
        limited = normal_eye()
        limited["quality"] = "LIMITED"
        limited["_source_filename"] = "ancillary.png"
        merged = app.merge_extractions([
            {"eyes": [adequate], "global_warnings": []},
            {"eyes": [limited], "global_warnings": []},
        ])
        eye = merged["eyes"][0]
        assert eye["quality"] == "ADEQUATE"
        assert not any("source image quality" in item for item in eye["data_conflicts"])

        merged_late_adequate = app.merge_extractions([
            {"eyes": [limited], "global_warnings": []},
            {"eyes": [dict(limited, _source_filename="second-limited.png")], "global_warnings": []},
            {"eyes": [adequate], "global_warnings": []},
        ])
        late_eye = merged_late_adequate["eyes"][0]
        assert late_eye["quality"] == "ADEQUATE"
        assert not any("source image quality" in item for item in late_eye["data_conflicts"])

        all_limited = app.merge_extractions([
            {"eyes": [limited], "global_warnings": []},
            {"eyes": [dict(limited, _source_filename="second-limited.png")], "global_warnings": []},
        ])
        limited_eye = all_limited["eyes"][0]
        assert limited_eye["quality"] == "LIMITED"
        assert not any("source image quality" in item for item in limited_eye["data_conflicts"])
        limited_result = app.assess_eye(limited_eye, plan(), 35, MODIFIERS)
        assert limited_result["status"] == "PASS"
        assert "adequate-quality tomography/topography" not in limited_result["missing"]

    def test_unreadable_anterior_map_prohibits_pass(self):
        eye = normal_eye()
        eye["anterior_pattern"] = "UNREADABLE"
        result = app.assess_eye(eye, plan(), 35, MODIFIERS)
        self.assertEqual(result["status"], "DATA INSUFFICIENT")
        self.assertIn("readable anterior pattern", result["missing"])

    def test_added_patient_modifiers_are_reported_without_invented_score_weights(self):
        modifiers = dict(
            MODIFIERS,
            pregnancy_nursing="yes",
            collagen_tissue_disease="yes",
            drug_usage="yes",
        )
        result = app.assess_eye(normal_eye(), plan(), 35, modifiers)
        self.assertEqual(result["status"], "STOP-DEFER")
        self.assertTrue(any("Pregnancy or nursing" in item for item in result["clinical_modifiers"]))
        self.assertTrue(any("Collagen/connective-tissue" in item for item in result["clinical_modifiers"]))
        self.assertTrue(any("medication/drug" in item for item in result["clinical_modifiers"]))

    def test_conflicting_i_s_values_use_concerning_value_and_prohibit_pass(self):
        first = normal_eye()
        second = normal_eye()
        first["I_S"] = 0.5
        second["I_S"] = 1.6
        merged = app.merge_extractions(
            [{"eyes": [first], "global_warnings": []}, {"eyes": [second], "global_warnings": []}]
        )
        self.assertEqual(merged["eyes"][0]["I_S"], 1.6)
        self.assertTrue(any(item.startswith("I_S:") for item in merged["eyes"][0]["data_conflicts"]))
        result = app.assess_eye(merged["eyes"][0], plan(), 35, MODIFIERS)
        self.assertNotEqual(result["status"], "PASS")

    def test_minor_keratometry_ocr_variation_is_not_an_unresolved_conflict(self):
        first = normal_eye()
        second = normal_eye()
        first["K2_D"] = 46.8
        second["K2_D"] = 46.6
        merged = app.merge_extractions(
            [{"eyes": [first], "global_warnings": []}, {"eyes": [second], "global_warnings": []}]
        )
        self.assertEqual(merged["eyes"][0]["K2_D"], 46.8)
        self.assertFalse(any("K2_D" in item for item in merged["eyes"][0]["data_conflicts"]))
        result = app.assess_eye(merged["eyes"][0], plan(), 35, MODIFIERS)
        self.assertEqual(result["status"], "PASS")

    def test_unverified_map_spot_number_is_not_accepted_as_a_table_parameter(self):
        eye = normal_eye()
        eye["K2_D"] = 46.6
        eye["table_verified_numeric_fields"].remove("K2_D")
        merged = app.merge_extractions([{"eyes": [eye], "global_warnings": []}])
        extracted = merged["eyes"][0]
        self.assertIsNone(extracted["K2_D"])
        self.assertIn("K2_D", extracted["missing_or_unreadable"])
        self.assertEqual(extracted["morphology"], "NORMAL_SYMMETRIC")

    def test_labeled_side_table_number_is_retained(self):
        eye = normal_eye()
        eye["K2_D"] = 46.8
        merged = app.merge_extractions([{"eyes": [eye], "global_warnings": []}])
        self.assertEqual(merged["eyes"][0]["K2_D"], 46.8)

    def test_keratometry_is_accepted_only_from_show2_topometric_cornea_front(self):
        wrong_source = normal_eye()
        wrong_source.update(
            keratometry_source="OTHER_PENTACAM_SOURCE",
            K1_D=47.1,
            K1_axis_deg=12,
            K2_D=48.2,
            K2_axis_deg=102,
            Kmean_D=47.6,
        )
        accepted = normal_eye()
        accepted.update(
            keratometry_source="SHOW_2_EXAMS_TOPOMETRIC_CORNEA_FRONT",
            K1_D=42.1,
            K1_axis_deg=5,
            K2_D=42.7,
            K2_axis_deg=95,
            Kmean_D=42.4,
        )

        merged = app.merge_extractions([
            {"eyes": [wrong_source], "global_warnings": []},
            {"eyes": [accepted], "global_warnings": []},
        ])
        eye = merged["eyes"][0]

        self.assertEqual(eye["K1_D"], 42.1)
        self.assertEqual(eye["K1_axis_deg"], 5)
        self.assertEqual(eye["K2_D"], 42.7)
        self.assertEqual(eye["K2_axis_deg"], 95)
        self.assertEqual(eye["Kmean_D"], 42.4)
        self.assertEqual(
            eye["keratometry_source"],
            "SHOW_2_EXAMS_TOPOMETRIC_CORNEA_FRONT",
        )
        self.assertFalse(any(
            field in str(conflict)
            for conflict in eye["data_conflicts"]
            for field in ("K1_D", "K1_axis_deg", "K2_D", "K2_axis_deg", "Kmean_D")
        ))

    def test_noncanonical_keratometry_is_removed_when_no_show2_source_exists(self):
        eye = normal_eye()
        eye["keratometry_source"] = "OTHER_PENTACAM_SOURCE"
        extracted = app.merge_extractions([{"eyes": [eye], "global_warnings": []}])["eyes"][0]

        for field in ("K1_D", "K1_axis_deg", "K2_D", "K2_axis_deg", "Kmean_D"):
            self.assertIsNone(extracted[field])
            self.assertNotIn(field, extracted["table_verified_numeric_fields"])
            self.assertIn(field, extracted["missing_or_unreadable"])

    def test_thinnest_map_number_is_rejected_when_labeled_row_is_unreadable(self):
        eye = normal_eye(pachy=566)
        eye["table_verified_numeric_fields"].remove("pachy_thinnest_um")
        eye["map_fallback_numeric_fields"] = ["pachy_thinnest_um"]
        merged = app.merge_extractions([{"eyes": [eye], "global_warnings": []}])
        extracted = merged["eyes"][0]
        self.assertIsNone(extracted["pachy_thinnest_um"])
        self.assertEqual(extracted["map_fallback_numeric_fields"], [])
        self.assertIn("pachy_thinnest_um", extracted["missing_or_unreadable"])

    def test_map_spot_cannot_substitute_for_k2_even_when_model_labels_it_as_fallback(self):
        eye = normal_eye()
        eye["K2_D"] = 46.6
        eye["table_verified_numeric_fields"].remove("K2_D")
        eye["map_fallback_numeric_fields"] = ["K2_D"]
        merged = app.merge_extractions([{"eyes": [eye], "global_warnings": []}])
        extracted = merged["eyes"][0]
        self.assertIsNone(extracted["K2_D"])
        self.assertNotIn("K2_D", extracted["map_fallback_numeric_fields"])

    def test_local_kmax_is_rejected_but_explicit_local_rmin_remains_permitted(self):
        eye = normal_eye()
        for field in ("Kmax_D", "Rmin_mm"):
            eye["table_verified_numeric_fields"].remove(field)
        eye["Kmax_D"] = 47.6
        eye["Rmin_mm"] = 7.09
        eye["map_fallback_numeric_fields"] = ["Kmax_D", "Rmin_mm"]
        extracted = app.merge_extractions([{"eyes": [eye], "global_warnings": []}])["eyes"][0]
        self.assertIsNone(extracted["Kmax_D"])
        self.assertEqual(extracted["Rmin_mm"], 7.09)
        self.assertEqual(extracted["map_fallback_numeric_fields"], ["Rmin_mm"])

    def test_multiple_local_rmin_values_do_not_create_unresolved_conflict(self):
        first = normal_eye()
        second = normal_eye()
        for eye in (first, second):
            for field in ("Kmax_D", "Rmin_mm"):
                eye["table_verified_numeric_fields"].remove(field)
            eye["map_fallback_numeric_fields"] = ["Kmax_D", "Rmin_mm"]
        first.update(Kmax_D=47.6, Rmin_mm=7.09)
        second.update(Kmax_D=48.1, Rmin_mm=7.03)
        merged = app.merge_extractions(
            [{"eyes": [first], "global_warnings": []}, {"eyes": [second], "global_warnings": []}]
        )
        extracted = merged["eyes"][0]
        self.assertIsNone(extracted["Kmax_D"])
        self.assertEqual(extracted["Rmin_mm"], 7.03)
        self.assertFalse(any("Rmin_mm" in item for item in extracted["data_conflicts"]))

    def test_labeled_edge_box_overrides_explicit_local_kmax_without_conflict(self):
        local = normal_eye()
        local["table_verified_numeric_fields"].remove("Kmax_D")
        local["Kmax_D"] = 48.1
        local["map_fallback_numeric_fields"] = ["Kmax_D"]
        table = normal_eye()
        table["Kmax_D"] = 47.6
        merged = app.merge_extractions(
            [{"eyes": [local], "global_warnings": []}, {"eyes": [table], "global_warnings": []}]
        )
        extracted = merged["eyes"][0]
        self.assertEqual(extracted["Kmax_D"], 47.6)
        self.assertNotIn("Kmax_D", extracted["map_fallback_numeric_fields"])
        self.assertFalse(any("Kmax_D" in item for item in extracted["data_conflicts"]))

    def test_later_labeled_table_value_overrides_earlier_local_map_fallback_without_conflict(self):
        map_eye = normal_eye(pachy=565)
        map_eye["table_verified_numeric_fields"].remove("pachy_thinnest_um")
        map_eye["map_fallback_numeric_fields"] = ["pachy_thinnest_um"]
        table_eye = normal_eye(pachy=566)
        merged = app.merge_extractions(
            [{"eyes": [map_eye], "global_warnings": []}, {"eyes": [table_eye], "global_warnings": []}]
        )
        extracted = merged["eyes"][0]
        self.assertEqual(extracted["pachy_thinnest_um"], 566)
        self.assertNotIn("pachy_thinnest_um", extracted["map_fallback_numeric_fields"])
        self.assertFalse(any("pachy_thinnest_um" in item for item in extracted["data_conflicts"]))

    def test_table_verified_field_lists_are_unioned_without_conflict(self):
        first = normal_eye()
        second = normal_eye()
        first["table_verified_numeric_fields"] = ["K1_D"]
        second["table_verified_numeric_fields"] = ["K2_D"]
        merged = app.merge_extractions(
            [{"eyes": [first], "global_warnings": []}, {"eyes": [second], "global_warnings": []}]
        )
        eye = merged["eyes"][0]
        self.assertEqual(eye["table_verified_numeric_fields"], ["K1_D", "K2_D"])
        self.assertEqual(eye["K1_D"], 42.0)
        self.assertEqual(eye["K2_D"], 43.0)
        self.assertFalse(any("table_verified" in item for item in eye["data_conflicts"]))

    def test_duplicate_cornea_front_keratometry_retains_first_without_conflict(self):
        first = normal_eye()
        second = normal_eye()
        second["K2_D"] = 43.3
        merged = app.merge_extractions(
            [{"eyes": [first], "global_warnings": []}, {"eyes": [second], "global_warnings": []}]
        )
        self.assertEqual(merged["eyes"][0]["K2_D"], 43.0)
        self.assertFalse(any("K2_D" in item for item in merged["eyes"][0]["data_conflicts"]))

    def test_uncertain_or_unreadable_page_does_not_conflict_with_readable_page(self):
        readable = normal_eye()
        limited = normal_eye()
        limited.update(
            morphology="UNCERTAIN",
            asymmetric_bow_tie="UNCERTAIN",
            srax="UNCERTAIN",
            anterior_pattern="UNREADABLE",
            posterior_pattern="UNREADABLE",
        )
        merged = app.merge_extractions(
            [{"eyes": [readable], "global_warnings": []}, {"eyes": [limited], "global_warnings": []}]
        )
        eye = merged["eyes"][0]
        self.assertEqual(eye["morphology"], "NORMAL_SYMMETRIC")
        self.assertEqual(eye["posterior_pattern"], "REASSURING")
        self.assertEqual(eye["data_conflicts"], [])

    def test_legacy_morphology_confidence_conflict_never_requires_surgeon_input(self):
        eye = normal_eye()
        eye["data_conflicts"] = ["morphology_confidence: UNREADABLE vs HIGH"]
        self.assertFalse(any(
            "morphology_confidence" in item for item in app.required_tomography_missing(eye)
        ))

    def test_unsupported_asymmetric_bowtie_label_does_not_override_supported_normal_map(self):
        normal = normal_eye()
        unsupported = normal_eye(morphology="ASYMMETRIC_BOWTIE")
        unsupported["asymmetric_bow_tie"] = "YES"
        unsupported["inferior_opposite_steepening_D"] = None
        merged = app.merge_extractions(
            [{"eyes": [normal], "global_warnings": []}, {"eyes": [unsupported], "global_warnings": []}]
        )
        eye = merged["eyes"][0]
        self.assertEqual(eye["morphology"], "NORMAL_SYMMETRIC")
        self.assertFalse(any("morphology" in item for item in eye["data_conflicts"]))

    def test_bilateral_engine_keeps_eyes_separate_and_overall_is_worst(self):
        extracted = {"eyes": [normal_eye("OS", 479), normal_eye("OD", 560)], "global_warnings": []}
        modifiers = dict(MODIFIERS, assessed_eyes=["OD", "OS"])
        decision = app.hc_engine(
            extracted, 35, {"OD": plan(), "OS": plan()}, modifiers
        )
        self.assertEqual([r["eye"] for r in decision["eyes"]], ["OD", "OS"])
        self.assertEqual(decision["eyes"][0]["status"], "PASS")
        self.assertEqual(decision["eyes"][1]["status"], "STOP-DEFER")
        self.assertEqual(decision["status"], "STOP-DEFER")

    def test_merge_clears_missing_when_later_image_supplies_value(self):
        first = normal_eye()
        first["BAD_D"] = None
        first["missing_or_unreadable"] = ["BAD_D"]
        second = normal_eye()
        merged = app.merge_extractions(
            [{"eyes": [first], "global_warnings": []}, {"eyes": [second], "global_warnings": []}]
        )
        self.assertEqual(merged["eyes"][0]["BAD_D"], 1.0)
        self.assertNotIn("BAD_D", merged["eyes"][0]["missing_or_unreadable"])


class TestTreatmentCardTransfer(unittest.TestCase):
    def test_surgeon_manifest_defaults_blank_intended_but_never_overwrites_it(self):
        entered = plan(sphere=None, cylinder=None)
        entered.update({
            "manifest_entered_sphere_D": -3.25,
            "manifest_cylinder_signed_D": -1.5,
            "entered_axis_deg": 80,
        })
        effective = app.apply_extracted_corrections(
            {"treatment_corrections": []}, {"OD": entered}
        )["OD"]
        self.assertEqual(effective["intended_entered_sphere_D"], -3.25)
        self.assertEqual(effective["intended_cylinder_signed_D"], -1.5)
        self.assertEqual(effective["intended_sphere_D"], -3.25)
        self.assertEqual(effective["intended_cylinder_magnitude_D"], 1.5)
        self.assertEqual(effective["correction_axis_deg"], 80)
        self.assertEqual(effective["intended_default_source"], "SURGEON_MANIFEST")

        entered.update({
            "intended_entered_sphere_D": -2.75,
            "intended_cylinder_signed_D": -1.0,
        })
        changed = app.apply_extracted_corrections(
            {"treatment_corrections": []}, {"OD": entered}
        )["OD"]
        self.assertEqual(changed["intended_entered_sphere_D"], -2.75)
        self.assertEqual(changed["intended_cylinder_signed_D"], -1.0)
        self.assertNotIn("intended_default_source", changed)

    def test_confident_card_fills_empty_plan_and_drives_zone_specific_ablation(self):
        extracted = {
            "eyes": [normal_eye()],
            "treatment_corrections": [card_correction()],
            "global_warnings": [],
        }
        empty = plan(sphere=None, cylinder=None, ablation=None)
        empty["optical_zone_mm"] = 6.5
        effective = app.apply_extracted_corrections(extracted, {"OD": empty})
        self.assertEqual(effective["OD"]["manifest_sphere_D"], -4.5)
        self.assertEqual(effective["OD"]["manifest_cylinder_magnitude_D"], 3.5)
        self.assertEqual(effective["OD"]["intended_sphere_D"], -4.5)
        self.assertEqual(effective["OD"]["intended_cylinder_magnitude_D"], 3.5)
        self.assertEqual(effective["OD"]["correction_axis_deg"], 170.0)
        self.assertIn("Duzeltme Miktari", effective["OD"]["correction_source"])
        result = app.assess_eye(normal_eye(), effective["OD"], 35, MODIFIERS)
        self.assertEqual(result["values"]["max_ablation_um"], 120.0)

    def test_manual_pair_wins_when_card_differs(self):
        extracted = {"treatment_corrections": [card_correction()]}
        effective = app.apply_extracted_corrections(extracted, {"OD": plan(sphere=-2, cylinder=1)})
        self.assertEqual(effective["OD"]["intended_sphere_D"], -2)
        self.assertEqual(effective["OD"]["intended_cylinder_magnitude_D"], 1)
        self.assertNotIn("correction_source", effective["OD"])
        self.assertTrue(any("manual intended correction differs" in item for item in effective["OD"]["correction_warnings"]))

    def test_role_specific_manual_value_only_overrides_that_role(self):
        entered = plan(sphere=None, cylinder=None)
        entered["intended_sphere_D"] = -2.0
        entered["intended_cylinder_magnitude_D"] = 1.0
        effective = app.apply_extracted_corrections(
            {"treatment_corrections": [card_correction()]}, {"OD": entered}
        )
        self.assertEqual(effective["OD"]["manifest_sphere_D"], -4.5)
        self.assertEqual(effective["OD"]["manifest_cylinder_magnitude_D"], 3.5)
        self.assertEqual(effective["OD"]["intended_sphere_D"], -2.0)
        self.assertEqual(effective["OD"]["intended_cylinder_magnitude_D"], 1.0)
        self.assertIn("(manifest)", effective["OD"]["correction_source"])

    def test_partial_manual_pair_is_never_mixed_with_card(self):
        partial = plan(sphere=-2, cylinder=None)
        effective = app.apply_extracted_corrections(
            {"treatment_corrections": [card_correction()]}, {"OD": partial}
        )
        self.assertEqual(effective["OD"]["intended_sphere_D"], -2)
        self.assertIsNone(effective["OD"]["intended_cylinder_magnitude_D"])
        self.assertTrue(any("partial manual intended correction" in item for item in effective["OD"]["correction_warnings"]))

    def test_uncertain_axis_keeps_confident_sphere_and_cylinder_only(self):
        correction = card_correction(axis=None, axis_status="UNCERTAIN")
        effective = app.apply_extracted_corrections(
            {"treatment_corrections": [correction]},
            {"OD": plan(sphere=None, cylinder=None)},
        )
        self.assertEqual(effective["OD"]["intended_sphere_D"], -4.5)
        self.assertEqual(effective["OD"]["intended_cylinder_magnitude_D"], 3.5)
        self.assertNotIn("correction_axis_deg", effective["OD"])
        self.assertTrue(any("axis was not confidently readable" in item for item in effective["OD"]["correction_warnings"]))

    def test_uncertain_or_conflicting_values_do_not_auto_fill(self):
        uncertain = card_correction(
            sphere=None, cylinder=None, sphere_cylinder_status="UNCERTAIN"
        )
        conflicting = card_correction(sphere=-3.5, cylinder=-3.0, axis=3)
        empty = plan(sphere=None, cylinder=None)
        uncertain_result = app.apply_extracted_corrections(
            {"treatment_corrections": [uncertain]}, {"OD": empty}
        )
        conflict_result = app.apply_extracted_corrections(
            {"treatment_corrections": [card_correction(), conflicting]}, {"OD": empty}
        )
        self.assertIsNone(uncertain_result["OD"]["intended_sphere_D"])
        self.assertIsNone(conflict_result["OD"]["intended_sphere_D"])
        self.assertTrue(any("Conflicting" in item for item in conflict_result["OD"]["correction_warnings"]))

    def test_plus_cylinder_is_not_transposed_or_auto_filled(self):
        plus = card_correction(cylinder=3.5)
        effective = app.apply_extracted_corrections(
            {"treatment_corrections": [plus]},
            {"OD": plan(sphere=None, cylinder=None)},
        )
        self.assertIsNone(effective["OD"]["intended_sphere_D"])
        self.assertIsNone(effective["OD"]["intended_cylinder_magnitude_D"])
        self.assertTrue(any("plus-cylinder" in item for item in effective["OD"]["correction_warnings"]))

    def test_unknown_eye_from_card_image_is_not_assessed(self):
        unknown = normal_eye("UNKNOWN")
        extracted = {
            "eyes": [normal_eye(), unknown],
            "treatment_corrections": [card_correction()],
            "global_warnings": [],
        }
        effective = app.apply_extracted_corrections(
            extracted, {"OD": plan(sphere=None, cylinder=None)}
        )
        decision = app.hc_engine(extracted, 35, effective, MODIFIERS)
        self.assertEqual([eye["eye"] for eye in decision["eyes"]], ["OD"])

    def test_merge_keeps_and_deduplicates_card_readings(self):
        correction = card_correction()
        merged = app.merge_extractions([
            {"eyes": [normal_eye()], "treatment_corrections": [correction], "global_warnings": []},
            {"eyes": [], "treatment_corrections": [correction], "global_warnings": []},
        ])
        self.assertEqual(merged["treatment_corrections"], [correction])


class TestApiIntegration(unittest.TestCase):
    def test_analyze_endpoint_accepts_eye_specific_payload(self):
        extraction = {
            "eyes": [normal_eye()],
            "treatment_corrections": [card_correction()],
            "global_warnings": [],
        }
        fake_response = SimpleNamespace(
            output_text=json.dumps(extraction), status="completed", incomplete_details=None
        )
        fake_client = SimpleNamespace(
            responses=SimpleNamespace(create=lambda **kwargs: fake_response)
        )
        client = TestClient(app.app)
        with patch.object(app, "openai_client", return_value=fake_client):
            response = client.post(
                "/analyze",
                files={"images": ("od.png", b"synthetic-image-bytes", "image/png")},
                data={
                    "age": "35",
                    "eye_plans": json.dumps({"OD": plan(sphere=None, cylinder=None, ablation=None)}),
                    "patient_modifiers": json.dumps(MODIFIERS),
                },
            )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["workflow_status"], "NEEDS_INPUT")
        self.assertNotIn("decision", payload)
        self.assertTrue(any("Both OD and OS" in item["message"] for item in payload["missing"]))
        self.assertEqual(payload["effective_eye_plans"]["OD"]["manifest_sphere_D"], -4.5)
        self.assertEqual(payload["effective_eye_plans"]["OD"]["intended_sphere_D"], -4.5)
        self.assertEqual(payload["effective_eye_plans"]["OD"]["correction_axis_deg"], 170.0)

    def test_analyze_retry_reuses_same_inflight_or_completed_assessment(self):
        extraction = {
            "eyes": [normal_eye()],
            "treatment_corrections": [card_correction()],
            "global_warnings": [],
        }
        fake_response = SimpleNamespace(
            output_text=json.dumps(extraction), status="completed", incomplete_details=None
        )
        create = Mock(return_value=fake_response)
        fake_client = SimpleNamespace(responses=SimpleNamespace(create=create))
        client = TestClient(app.app)
        request_id = "c5946726-b4fd-4dd8-b47d-e8b44d00cf12"
        request = {
            "files": {"images": ("od.png", b"synthetic-image-bytes", "image/png")},
            "data": {
                "age": "35",
                "eye_plans": json.dumps({"OD": plan(sphere=None, cylinder=None, ablation=None)}),
                "patient_modifiers": json.dumps(MODIFIERS),
                "assessment_request_id": request_id,
            },
        }
        with patch.object(app, "openai_client", return_value=fake_client):
            first = client.post("/analyze", **request)
            calls_after_first_request = create.call_count
            second = client.post("/analyze", **request)
        self.assertEqual(first.status_code, 200)
        self.assertEqual(second.status_code, 200)
        self.assertEqual(first.json()["assessment_token"], second.json()["assessment_token"])
        self.assertGreater(calls_after_first_request, 0)
        self.assertEqual(create.call_count, calls_after_first_request)

    def test_professional_pdf_and_word_exports_are_valid(self):
        extracted = {"eyes": [normal_eye(), normal_eye("OS")], "global_warnings": []}
        decision = app.hc_engine(extracted, 35, {"OD": plan()}, MODIFIERS)
        payload = {
            "patient": {
                "name": "Test Patient", "id": "CER-AI-001", "age": 35,
                "reviewer": "Test Reviewer", "report_date": "2026-08-25",
            },
            "decision": decision,
            "extracted": extracted,
        }
        client = TestClient(app.app)
        from assessment_workflow import begin
        ready = begin(app, extracted, 35, {"OD": plan(), "OS": plan()}, MODIFIERS, payload["patient"])
        self.assertEqual(ready["workflow_status"], "READY")
        payload.update({"assessment_token": ready["assessment_token"], "report_token": ready["report_token"]})
        pdf = client.post("/report/pdf", json=payload)
        word = client.post("/report/word", json=payload)
        self.assertEqual(pdf.status_code, 200)
        self.assertEqual(pdf.headers["content-type"], "application/pdf")
        pdf_reader = PdfReader(BytesIO(pdf.content))
        self.assertGreaterEqual(len(pdf_reader.pages), 1)
        pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
        self.assertIn("Cornea Ectasia Risk Assessment Intelligence", pdf_text)
        self.assertIn(
            "The final surgical decision and all associated responsibility and liability rest with the surgeon. "
            "This application is a clinical decision-support aid only.",
            " ".join(pdf_text.split()),
        )
        self.assertGreaterEqual(pdf_text.count("TEST PATIENT"), 1)
        self.assertLess(pdf_text.rindex("TEST PATIENT"), pdf_text.index("OVERALL DISPOSITION"))
        self.assertEqual(word.status_code, 200)
        self.assertIn("wordprocessingml.document", word.headers["content-type"])
        document = Document(BytesIO(word.content))
        self.assertIn(
            "The final surgical decision and all associated responsibility and liability rest with the surgeon. "
            "This application is a clinical decision-support aid only.",
            "\n".join(paragraph.text for paragraph in document.paragraphs),
        )
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Cornea Ectasia Risk Assessment Intelligence", text)
        self.assertEqual(document.core_properties.title, "CER-AI — Cornea Ectasia Risk Assessment Intelligence")
        self.assertIn("CER-AI PREOPERATIVE ECTASIA RISK ASSESSMENT", text)
        self.assertIn("PASS", text)
        patient_paragraph = next(p for p in document.paragraphs if p.text == "TEST PATIENT")
        self.assertTrue(patient_paragraph.runs[0].bold)
        self.assertEqual(patient_paragraph.runs[0].font.size, Pt(15))
        body_children = list(document.element.body)
        patient_index = next(
            index for index, child in enumerate(body_children)
            if "".join(node.text or "" for node in child.xpath(".//w:t")) == "TEST PATIENT"
        )
        next_text = "".join(
            node.text or "" for node in body_children[patient_index + 1].xpath(".//w:t")
        )
        self.assertIn("OVERALL DISPOSITION", next_text)

    def test_quality_warning_is_at_the_end_of_pdf_and_word_without_blocking_report(self):
        od = normal_eye()
        od.update(quality="LIMITED", pentacam_qs="NOT_OK")
        extracted = {"eyes": [od, normal_eye("OS")], "global_warnings": []}
        decision = app.hc_engine(extracted, 35, {"OD": plan(), "OS": plan()}, MODIFIERS)
        self.assertEqual(decision["status"], "PASS")
        payload = {"patient": {"name": "Quality Warning"}, "decision": decision, "extracted": extracted}

        pdf_text = "\n".join(
            page.extract_text() or ""
            for page in PdfReader(BytesIO(reports.build_pdf(payload))).pages
        )
        self.assertIn("PENTACAM ACQUISITION QUALITY", pdf_text)
        self.assertIn("QS is NOT_OK", pdf_text)
        self.assertGreater(pdf_text.index("PENTACAM ACQUISITION QUALITY"), pdf_text.index("Interpretation note"))

        document = Document(BytesIO(reports.build_docx(payload)))
        body_text = "\n".join(
            "".join(node.text or "" for node in child.xpath(".//w:t"))
            for child in document.element.body
        )
        self.assertIn("PENTACAM ACQUISITION QUALITY", body_text)
        self.assertIn("QS is NOT_OK", body_text)
        self.assertGreater(body_text.index("PENTACAM ACQUISITION QUALITY"), body_text.index("Interpretation note"))

    def test_browser_report_places_large_bold_patient_name_immediately_before_status(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        patient_banner = '<div id="reportPatientBanner" class="report-patient-name">'
        status = '<div id="status" class="status insufficient">'
        self.assertLess(html.index(patient_banner), html.index(status))
        between = html[html.index(patient_banner):html.index(status)]
        self.assertNotIn("<table", between)
        self.assertIn(".report-patient-name{font-size:24px", html)
        self.assertIn("font-weight:800", html)
        self.assertIn("text-transform:uppercase", html)
        self.assertIn('document.querySelector("#reportPatientBanner").textContent=p.name', html)

    def test_reports_always_render_od_before_os(self):
        modifiers = dict(MODIFIERS, assessed_eyes=["OD", "OS"])
        extracted = {"eyes": [normal_eye("OS"), normal_eye("OD")], "global_warnings": []}
        decision = app.hc_engine(extracted, 35, {"OD": plan(), "OS": plan()}, modifiers)
        self.assertEqual([eye["eye"] for eye in decision["eyes"]], ["OD", "OS"])

        reversed_payload = {
            "patient": {},
            "decision": dict(decision, eyes=list(reversed(decision["eyes"]))),
            "extracted": extracted,
        }
        pdf_text = "\n".join(
            page.extract_text() or ""
            for page in PdfReader(BytesIO(reports.build_pdf(reversed_payload))).pages
        )
        self.assertLess(pdf_text.index("OD ASSESSMENT"), pdf_text.index("OS ASSESSMENT"))

        document = Document(BytesIO(reports.build_docx(reversed_payload)))
        word_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertLess(word_text.index("OD assessment"), word_text.index("OS assessment"))

        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn("orderedEyes(j.decision.eyes).map", html)


class TestPwaIcons(unittest.TestCase):
    def test_icon_files_are_real_pngs_with_declared_dimensions(self):
        expected = {
            "icon-192.png": (192, 192),
            "icon-512.png": (512, 512),
            "icon-maskable-512.png": (512, 512),
        }
        icon_dir = Path(__file__).resolve().parents[1] / "static" / "icons"
        for filename, dimensions in expected.items():
            raw = (icon_dir / filename).read_bytes()
            self.assertEqual(raw[:8], b"\x89PNG\r\n\x1a\n")
            self.assertEqual(raw[12:16], b"IHDR")
            self.assertEqual(struct.unpack(">II", raw[16:24]), dimensions)

    def test_manifest_and_html_use_cache_busted_png_assets(self):
        static_dir = Path(__file__).resolve().parents[1] / "static"
        manifest = json.loads((static_dir / "manifest.webmanifest").read_text())
        self.assertEqual(manifest["name"], "CER-AI — Cornea Ectasia Risk Assessment Intelligence")
        self.assertEqual({icon["src"] for icon in manifest["icons"]}, {
            "/static/icons/icon-192.png?v=8",
            "/static/icons/icon-512.png?v=8",
            "/static/icons/icon-maskable-512.png?v=8",
        })
        html = (static_dir / "index.html").read_text()
        self.assertIn('/static/manifest.webmanifest?v=9', html)
        self.assertIn('/static/icons/favicon-32.png?v=8', html)
        self.assertIn('/static/icons/apple-touch-icon.png?v=8', html)
        self.assertIn('/static/branding/cer-ai-logo-final.png?v=4', html)
        self.assertIn('alt="CER-AI — Cornea Ectasia Risk Assessment Intelligence"', html)
        self.assertIn('Cornea Ectasia Risk <span class="brand-initial">A</span>ssessment <span class="brand-initial">I</span>ntelligence', html)
        self.assertIn('.brand-subtitle .brand-initial{color:#ef7f2d}', html)
        self.assertNotIn("Corneal Ectasia Risk Analysis Intelligence", html)
        self.assertNotIn('/static/icons/icon-source.svg', html)


class TestAuthorshipAndLiabilityFooter(unittest.TestCase):
    NOTICE = (
        "Developed by Hüseyin Cengiz, MD. All rights reserved. "
        "Final responsibility rests with the surgeon at all times and under all circumstances."
    )

    def _payload(self):
        extracted = {"eyes": [normal_eye(), normal_eye("OS")], "global_warnings": []}
        return {
            "patient": {"name": "Footer Test", "age": 35},
            "decision": app.hc_engine(extracted, 35, {"OD": plan(), "OS": plan()}, MODIFIERS),
            "extracted": extracted,
        }

    def test_interface_and_print_report_contain_exact_notice(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertEqual(html.count(self.NOTICE), 2)
        self.assertIn('class="app-legal-footer"', html)
        self.assertIn('class="report-legal-footer"', html)
        self.assertIn(".report-legal-footer{position:fixed", html)

    def test_pdf_notice_is_present_on_every_page(self):
        reader = PdfReader(BytesIO(reports.build_pdf(self._payload())))
        self.assertGreaterEqual(len(reader.pages), 2)
        for page in reader.pages:
            self.assertIn(self.NOTICE, " ".join((page.extract_text() or "").split()))

    def test_word_notice_is_in_every_section_footer(self):
        document = Document(BytesIO(reports.build_docx(self._payload())))
        self.assertTrue(document.sections)
        for section in document.sections:
            footer_text = " ".join(paragraph.text for paragraph in section.footer.paragraphs)
            self.assertIn(self.NOTICE, footer_text)

    def test_pdf_and_word_contain_the_same_simplified_topography_reference(self):
        pdf_text = " ".join(
            " ".join((page.extract_text() or "").split())
            for page in PdfReader(BytesIO(reports.build_pdf(self._payload()))).pages
        )
        document = Document(BytesIO(reports.build_docx(self._payload())))
        word_text = " ".join(paragraph.text for paragraph in document.paragraphs)
        word_text += " " + " ".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        for text in (pdf_text, word_text):
            self.assertIn("Randleman topography assessment", text)
            self.assertIn("Mild asymmetric bow-tie", text)
            self.assertIn("Only the highest applicable single category is scored", text)
            self.assertIn("Superior steepening alone is not automatically assigned 3 points", text)

    def test_turkish_pdf_uses_unicode_labels_and_turkish_footer_on_every_page(self):
        payload = self._payload()
        payload["locale"] = "tr"
        reader = PdfReader(BytesIO(reports.build_pdf(payload)))
        pages = [" ".join((page.extract_text() or "").split()) for page in reader.pages]
        combined = " ".join(pages)
        self.assertIn("CER-AI PREOPERATİF EKTAZİ RİSK DEĞERLENDİRMESİ", combined)
        self.assertIn("GENEL KARAR", combined)
        self.assertIn("UYGUN", combined)
        self.assertIn("Randleman topografi değerlendirmesi", combined)
        self.assertIn("≥21 yaş", combined)
        self.assertIn("Kesin durdurma / 2 / 1 / 0", combined)
        self.assertNotIn(">=21 years", combined)
        self.assertNotIn("Hard stop / 2 / 1 / 0", combined)
        for page in pages:
            self.assertIn("Hüseyin Cengiz, MD tarafından geliştirilmiştir", page)

    def test_turkish_word_uses_turkish_labels_and_footer(self):
        payload = self._payload()
        payload["locale"] = "tr"
        document = Document(BytesIO(reports.build_docx(payload)))
        body = " ".join(paragraph.text for paragraph in document.paragraphs)
        table_text = " ".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn("CER-AI PREOPERATİF EKTAZİ RİSK DEĞERLENDİRMESİ", body)
        self.assertIn("GENEL KARAR", table_text)
        self.assertIn("≥21 yaş", table_text)
        self.assertIn("Kesin durdurma / 2 / 1 / 0", table_text)
        self.assertNotIn(">=21 years", table_text)
        self.assertNotIn("Hard stop / 2 / 1 / 0", table_text)
        for section in document.sections:
            footer_text = " ".join(paragraph.text for paragraph in section.footer.paragraphs)
            self.assertIn("Hüseyin Cengiz, MD tarafından geliştirilmiştir", footer_text)

    def test_interface_has_persistent_english_turkish_selector_and_localized_exports(self):
        root = Path(__file__).resolve().parents[1]
        html = (root / "static" / "index.html").read_text()
        i18n = (root / "static" / "i18n.js").read_text()
        self.assertIn('data-language="en"', html)
        self.assertIn('data-language="tr"', html)
        self.assertIn('/static/i18n.js?v=5', html)
        self.assertIn('locale:i18n.locale', html)
        self.assertIn('localStorage.setItem("cerai-language"', i18n)
        self.assertIn('"Case inputs":"Vaka girdileri"', i18n)
        self.assertIn('"18 / 19–20 / ≥21 years":"18 / 19–20 / ≥21 yaş"', i18n)
        self.assertIn('"Hard stop / 2 / 1 / 0":"Kesin durdurma / 2 / 1 / 0"', i18n)


class TestSurgeonTopographyReference(unittest.TestCase):
    def test_randleman_confirmation_includes_source_locked_reference(self):
        root = Path(__file__).resolve().parents[1]
        html = (root / "static" / "index.html").read_text()
        reference = (root / "static" / "hc-reference.js").read_text()
        self.assertEqual(html.count('data-erss-reference="surgeon"'), 1)
        self.assertIn('aria-label="Randleman anterior topography reference guide"', reference)
        self.assertIn("upper-left Axial/Sagittal Curvature (Front)", reference)
        self.assertIn("It never guesses a number", reference)
        self.assertIn("BAD-D and other tomography indices are not substituted", reference)

    def test_reference_matches_published_mutually_exclusive_thresholds(self):
        reference = (Path(__file__).resolve().parents[1] / "static" / "hc-reference.js").read_text()
        for threshold in (">0.5 D and <1.0 D", "SRAX ≥20°", "≥1.0 D", "I-S <1.4 D", "I-S ≥1.4 D"):
            self.assertIn(threshold, reference)
        for category, points in (
            ("Normal / symmetric", "0"),
            ("Asymmetric bow-tie", "1"),
            ("Inferior steepening / SRA", "3"),
            ("Abnormal / ectatic", "4"),
        ):
            self.assertRegex(reference, rf'\["{category}",\s*"[^"]+",\s*"{points}"\]')
        self.assertIn("Only the highest applicable single category is scored; categories are not added", reference)
        self.assertIn("Superior steepening alone is not automatically assigned 3 points", reference)

    def test_obsolete_fragmented_reference_copy_is_removed(self):
        root = Path(__file__).resolve().parents[1]
        html = (root / "static" / "index.html").read_text()
        reference = (root / "static" / "hc-reference.js").read_text()
        self.assertNotIn("Surgeon reference — how to choose", html)
        self.assertNotIn("Published Randleman / ERSS scoring points", reference)
        self.assertNotIn('["Age", "18–21", "3"]', reference)


class TestSignedRefractionInputs(unittest.TestCase):
    def test_all_refraction_fields_accept_positive_and_negative_values(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('id="${eye}_manifest_cylinder" type="text" inputmode="decimal"', html)
        self.assertIn('id="${eye}_cylinder" type="text" inputmode="decimal"', html)
        self.assertIn('manifest_cylinder_signed_D:numberOrNull', html)
        self.assertIn('intended_cylinder_signed_D:numberOrNull', html)

    def test_all_refraction_fields_have_phone_safe_negative_buttons(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        for target in ("manifest_sphere", "manifest_cylinder", "sphere", "cylinder"):
            self.assertIn(f'data-negative-target="${{eye}}_{target}"', html)
        self.assertIn('applySign(input,"-")', html)

    def test_all_refraction_fields_have_phone_safe_positive_buttons(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        for target in ("manifest_sphere", "manifest_cylinder", "sphere", "cylinder"):
            self.assertIn(f'data-positive-target="${{eye}}_{target}"', html)
        self.assertIn('applySign(input,"+")', html)

    def test_empty_sign_button_inserts_only_sign_and_sign_only_is_not_numeric(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('input.value=unsigned?`${sign}${unsigned}`:sign;', html)
        self.assertIn('if(raw===""||raw==="+"||raw==="-")return {state:"blank",value:null};', html)
        self.assertNotIn('parsed||0.25', html)

    def test_mobile_minus_variants_are_normalized_before_numeric_parsing(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('replace(/[−–—﹣－]/g,"-")', html)
        self.assertIn('replace(/[＋﹢]/g,"+")', html)

    def test_invalid_or_partial_manifest_input_cannot_silently_become_missing(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('function validateRefractionInputs()', html)
        self.assertIn('Complete both ${eye.toUpperCase()} ${role} sphere and cylinder, or leave both blank.', html)
        self.assertIn('const invalidRefraction=validateRefractionInputs();', html)

    def test_autofilled_internal_cylinder_magnitudes_render_as_negative(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('cylinder.value=-Math.abs(p[`${role}_cylinder_magnitude_D`])', html)
        self.assertIn('!sphere.value.trim()&&!cylinder.value.trim()', html)

    def test_plus_cylinder_is_transposed_to_equivalent_minus_cylinder(self):
        entered = plan()
        entered.pop("manifest_sphere_D")
        entered.pop("manifest_cylinder_magnitude_D")
        entered.pop("intended_sphere_D")
        entered.pop("intended_cylinder_magnitude_D")
        entered.update({
            "manifest_entered_sphere_D": -5.0,
            "manifest_cylinder_signed_D": 2.0,
            "intended_entered_sphere_D": -5.0,
            "intended_cylinder_signed_D": 2.0,
            "entered_axis_deg": 90.0,
        })
        result = app.assess_eye(normal_eye(), entered, 35, MODIFIERS)
        values = result["values"]
        self.assertEqual(values["manifest_sphere_D"], -3.0)
        self.assertEqual(values["manifest_cylinder_magnitude_D"], 2.0)
        self.assertEqual(values["intended_sphere_D"], -3.0)
        self.assertEqual(values["correction_axis_deg"], 180.0)
        self.assertEqual(values["manifest_normalized_axis_deg"], 180.0)
        self.assertEqual(values["MRSE_D"], -4.0)
        self.assertEqual(values["manifest_entered_sphere_D"], -5.0)
        self.assertEqual(values["manifest_cylinder_signed_D"], 2.0)
        self.assertTrue(any("plus-cylinder notation was transposed" in item for item in result["warnings"]))

    def test_minus_cylinder_keeps_sphere_and_axis(self):
        normalized = app.normalize_signed_refraction_plan({
            "manifest_entered_sphere_D": 2.0,
            "manifest_cylinder_signed_D": -1.5,
            "intended_entered_sphere_D": 2.0,
            "intended_cylinder_signed_D": -1.5,
            "entered_axis_deg": 180.0,
        })
        self.assertEqual(normalized["manifest_sphere_D"], 2.0)
        self.assertEqual(normalized["manifest_cylinder_magnitude_D"], 1.5)
        self.assertEqual(normalized["intended_sphere_D"], 2.0)
        self.assertEqual(normalized["correction_axis_deg"], 180.0)

    def test_plus_cylinder_axis_rotation_wraps_at_180(self):
        self.assertEqual(app._transpose_axis(90), 180.0)
        self.assertEqual(app._transpose_axis(180), 90.0)

    def test_nonzero_cylinder_without_axis_prohibits_pass(self):
        entered = plan("LASIK", sphere=None, cylinder=None, ablation=55, flap=100)
        entered.update({
            "manifest_entered_sphere_D": -5.0,
            "manifest_cylinder_signed_D": 2.0,
            "intended_entered_sphere_D": -5.0,
            "intended_cylinder_signed_D": 2.0,
            "entered_axis_deg": None,
        })
        result = app.assess_eye(normal_eye(), entered, 35, MODIFIERS)
        self.assertEqual(result["status"], "DATA INSUFFICIENT")
        self.assertIn("manifest cylinder axis for non-zero cylinder", result["missing"])
        self.assertIn("intended cylinder axis for non-zero cylinder", result["missing"])

    def test_reports_disclose_entered_and_normalized_notation(self):
        source = (Path(__file__).resolve().parents[1] / "reports.py").read_text()
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        for label in (
            "Manifest entered notation", "Manifest normalized (minus-cylinder)",
            "Intended entered notation", "Intended normalized (minus-cylinder)",
        ):
            self.assertIn(label, source)
            self.assertIn(label, html)


class TestPatientIdentityWarningUi(unittest.TestCase):
    def test_unverified_identity_has_prominent_warning_without_hiding_eye_results(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('id="identityWarning" class="identity-alert" hidden', html)
        self.assertIn("PATIENT IDENTITY NOT VERIFIED — SURGEON CONFIRMATION REQUIRED", html)
        self.assertIn('orderedEyes(j.decision.eyes).map(x=>renderEye', html)


class TestPwaShareTarget(unittest.TestCase):
    def test_manifest_accepts_one_or_multiple_shared_images(self):
        static_dir = Path(__file__).resolve().parents[1] / "static"
        manifest = json.loads((static_dir / "manifest.webmanifest").read_text())
        target = manifest["share_target"]
        self.assertEqual(target["action"], "/share-target")
        self.assertEqual(target["method"], "POST")
        self.assertEqual(target["enctype"], "multipart/form-data")
        self.assertEqual(target["params"]["files"], [{"name": "images", "accept": ["image/*"]}])

    def test_root_scoped_service_worker_is_served_without_caching(self):
        response = TestClient(app.app).get("/sw.js")
        self.assertEqual(response.status_code, 200)
        self.assertIn("application/javascript", response.headers["content-type"])
        self.assertEqual(response.headers["service-worker-allowed"], "/")
        self.assertEqual(response.headers["cache-control"], "no-cache, no-store, must-revalidate")
        self.assertIn('url.pathname===SHARE_PATH', response.text)

    def test_frontend_loads_shared_files_into_existing_analyze_request(self):
        static_dir = Path(__file__).resolve().parents[1] / "static"
        html = (static_dir / "index.html").read_text()
        self.assertIn('id="imageInput" name="images"', html)
        self.assertIn('params.get("share_token")', html)
        self.assertIn('await appendSharedImages(fd,sharedImages)', html)
        self.assertIn('fd.append("images",blob,item.name', html)
        self.assertNotIn("new DataTransfer()", html)
        self.assertNotIn("new File([blob]", html)
        self.assertIn('navigator.serviceWorker.register("/sw.js",{scope:"/"})', html)

    def test_share_cache_retains_originals_for_preview_and_retry(self):
        static_dir = Path(__file__).resolve().parents[1] / "static"
        worker = (static_dir / "sw.js").read_text()
        html = (static_dir / "index.html").read_text()
        self.assertIn("size:file.size", worker)
        self.assertIn("url.pathname.startsWith(SHARE_STORAGE_PATH)", worker)
        self.assertIn("previewUrl:item.url", html)
        self.assertNotIn("await cache.delete(item.url)", html)
        self.assertNotIn("await cache.delete(metaUrl)", html)


class TestAnalyzeLoadingStateUi(unittest.TestCase):
    def test_loading_message_paints_before_network_request_and_validation_is_visible(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('<form id="f" novalidate>', html)
        self.assertIn('id="assessmentProgress"', html)
        self.assertIn('role="status" aria-live="polite"', html)
        self.assertIn('assessBtn.textContent="Assessing images... Please wait"', html)
        self.assertIn('showFormMessage("Assessing images... Please wait.")', html)
        self.assertIn("await allowLoadingStateToPaint()", html)
        self.assertLess(html.index("await allowLoadingStateToPaint()"), html.index('ceraiFetch("/analyze"'))
        self.assertIn("f.reportValidity()", html)
        self.assertIn('showFormMessage(`Assessment failed:', html)

    def test_mobile_fetch_failure_is_recovered_once_without_duplicate_assessment(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('fd.set("assessment_request_id",requestId)', html)
        self.assertIn('if(!isFetchConnectionError(error))throw error;', html)
        self.assertEqual(html.count('r=await sendAssessment();'), 2)
        self.assertIn("CER-AI is restoring this assessment automatically", html)


class TestPatientModifierUi(unittest.TestCase):
    def test_single_dropdown_contains_all_multi_select_modifier_options(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertEqual(html.count('id="modifierDropdown"'), 1)
        for value in (
            "eye_rubbing", "family_history", "inter_eye_asymmetry",
            "pregnancy_nursing", "collagen_tissue_disease", "drug_usage",
        ):
            self.assertIn(f'name="patient_modifier" value="{value}"', html)
        self.assertNotIn('id="eye_rubbing"', html)
        self.assertNotIn('id="family_history"', html)
        self.assertNotIn('id="inter_eye_asymmetry"', html)
        self.assertIn('value="none" data-exclusive="true"', html)
        self.assertIn('value="unknown" data-exclusive="true"', html)

    def test_none_closes_dropdown_but_regular_multi_select_items_do_not(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn('if(input.checked&&input.value==="none")modifierDropdown.open=false;', html)
        self.assertNotIn('if(input.checked)modifierDropdown.open=false;', html)


class TestFixedLaserPlatformUi(unittest.TestCase):
    def test_each_eye_keeps_a_read_only_alcon_ex500_box(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn(
            'id="${eye}_platform" type="text" value="Alcon EX500" readonly aria-readonly="true"',
            html,
        )
        self.assertIn('laser_platform:"Alcon EX500"', html)
        self.assertNotIn('placeholder="e.g., Alcon EX500"', html)


class TestLasikFlapDropdownUi(unittest.TestCase):
    def test_flap_field_is_a_dropdown_with_only_hc_options(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        expected = (
            '<select id="${eye}_flap"><option value="">Select</option>'
            '<option value="90">90 µm</option><option value="100" selected>100 µm</option>'
            '<option value="110">110 µm</option><option value="120">120 µm</option></select>'
        )
        self.assertIn(expected, html)
        self.assertNotIn('id="${eye}_flap" type="number"', html)


class TestFixedPrkEpitheliumUi(unittest.TestCase):
    def test_each_eye_has_a_read_only_50_micron_prk_epithelium_box(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn(
            'id="${eye}_epithelium" type="text" value="50" readonly aria-readonly="true"',
            html,
        )
        self.assertIn("epithelium_um:50", html)


class TestZoneDropdownUi(unittest.TestCase):
    def test_zone_fields_use_only_the_requested_dropdown_options(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        optical = (
            '<select id="${eye}_optical"><option value="">Select</option>'
            '<option value="6.0">6.0 mm</option><option value="6.5" selected>6.5 mm</option>'
            '<option value="7.0">7.0 mm</option></select>'
        )
        transition = (
            '<select id="${eye}_transition"><option value="">Select</option>'
            '<option value="8.0">8.0 mm</option><option value="8.5">8.5 mm</option>'
            '<option value="9.0" selected>9.0 mm</option></select>'
        )
        self.assertIn(optical, html)
        self.assertIn(transition, html)
        self.assertNotIn('id="${eye}_transition_na"', html)
        self.assertIn('transition_zone_not_applicable:"no"', html)


class TestPriorSurgeryDefaultUi(unittest.TestCase):
    def test_prior_surgery_defaults_to_no_and_yes_remains_selectable(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        expected = (
            '<select id="${eye}_prior"><option value="no" selected>No</option>'
            '<option value="yes">Yes</option><option value="unknown">Unknown</option></select>'
        )
        self.assertIn(expected, html)


class TestClinicalEligibilityGroupUi(unittest.TestCase):
    def test_four_eye_specific_clinical_controls_share_one_collapsible_box(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        self.assertIn(
            'id="${eye}_clinical" class="clinical-select-group">\n'
            '    <summary>Clinical eligibility and stability</summary>',
            html,
        )
        group_start = html.index('id="${eye}_clinical"')
        group_end = html.index('</details>`;', group_start)
        group = html[group_start:group_end]
        for field in ("stable", "progression", "cdva", "enhancement"):
            self.assertEqual(group.count(f'id="${{eye}}_{field}"'), 1)
        self.assertIn(
            'id="${eye}_stable"><option value="yes" selected>Yes</option>'
            '<option value="no">No</option><option value="unknown">Unknown</option>',
            group,
        )
        for field in ("progression", "cdva", "enhancement"):
            self.assertIn(
                f'id="${{eye}}_{field}"><option value="no" selected>No</option>'
                '<option value="yes">Yes</option><option value="unknown">Unknown</option>',
                group,
            )


class TestLiabilityNoticeUi(unittest.TestCase):
    def test_input_and_report_contain_the_same_surgeon_liability_notice(self):
        html = (Path(__file__).resolve().parents[1] / "static" / "index.html").read_text()
        notice = (
            '<p class="liability-notice">The final surgical decision and all associated responsibility and '
            'liability rest with the surgeon. This application is a clinical decision-support aid only.</p>'
        )
        self.assertIn(notice, html)
        self.assertIn(
            '<p class="report-liability-notice">The final surgical decision and all associated responsibility and '
            'liability rest with the surgeon. This application is a clinical decision-support aid only.</p>', html,
        )
        self.assertIn('.liability-notice{color:#a31212;', html)


if __name__ == "__main__":
    unittest.main()
