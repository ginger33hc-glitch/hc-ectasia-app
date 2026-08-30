"""Professional PDF and Word exports for CERAI reports."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List

import reportlab
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


NAVY = "173B57"
BLUE = "1F5E8C"
BLUE_FILL = "EAF3FA"
GREEN = "176B3A"
GREEN_FILL = "E6F4EA"
AMBER = "9A5A00"
AMBER_FILL = "FFF2DB"
RED = "A31212"
RED_FILL = "FDE8E8"
GRAY = "52616D"
GRAY_FILL = "EEF2F5"
LINE = "D7E0E7"
INK = "17212B"
LIABILITY_NOTICE = (
    "The final surgical decision and all associated responsibility and liability rest with the surgeon. "
    "This application is a clinical decision-support aid only."
)
PDF_UNICODE_BOLD = "CERAI-Vera-Bold"
pdfmetrics.registerFont(TTFont(
    PDF_UNICODE_BOLD,
    str(Path(reportlab.__file__).resolve().parent / "fonts" / "VeraBd.ttf"),
))


def _rl(value: str):
    return colors.HexColor(f"#{value}")


def _text(value: Any, fallback: str = "Not documented") -> str:
    if value is None or value == "":
        return fallback
    return str(value)


def _ascii(value: Any, fallback: str = "Not documented") -> str:
    return (
        _text(value, fallback)
        .replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2265", ">=")
        .replace("\u2264", "<=")
        .replace("\u00b5", "u")
        .replace("\u00b0", " degrees")
    )


def _status_palette(status: str) -> tuple[str, str]:
    if status == "PASS":
        return GREEN, GREEN_FILL
    if status == "DO NOT PROCEED":
        return RED, RED_FILL
    if status.startswith("CAUTION") or status.startswith("REVIEW"):
        return AMBER, AMBER_FILL
    return GRAY, GRAY_FILL


def _fmt(value: Any, digits: int = 1, unit: str = "") -> str:
    if value is None:
        return "Not available"
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        rendered = f"{value:.{digits}f}"
    else:
        rendered = str(value)
    return f"{rendered}{unit}"


def _eye_metrics(eye: Dict[str, Any]) -> List[tuple[str, str]]:
    values = eye.get("values") or {}
    score = eye.get("score") or {}
    erss_evidence = eye.get("erss_topography_evidence") or {}
    correction = "Not documented"
    if values.get("intended_sphere_D") is not None and values.get("intended_cylinder_magnitude_D") is not None:
        axis = (
            f" x {_fmt(values.get('correction_axis_deg'), 0, ' degrees')}"
            if values.get("correction_axis_deg") is not None else " (axis unavailable)"
        )
        correction = (
            f"{_fmt(values.get('intended_sphere_D'), 2, ' D')} / "
            f"-{_fmt(values.get('intended_cylinder_magnitude_D'), 2, ' D')}{axis}"
        )
    manifest = "Not documented"
    if values.get("manifest_sphere_D") is not None and values.get("manifest_cylinder_magnitude_D") is not None:
        manifest_axis = (
            f" x {_fmt(values.get('manifest_normalized_axis_deg'), 0, ' degrees')}"
            if values.get("manifest_normalized_axis_deg") is not None else " (axis unavailable)"
        )
        manifest = (
            f"{_fmt(values.get('manifest_sphere_D'), 2, ' D')} / "
            f"-{_fmt(values.get('manifest_cylinder_magnitude_D'), 2, ' D')}{manifest_axis}"
        )
    def entered_refraction(role: str) -> str:
        sphere = values.get(f"{role}_entered_sphere_D")
        cylinder = values.get(f"{role}_cylinder_signed_D")
        axis = values.get(f"{role}_entered_axis_deg")
        if sphere is None or cylinder is None:
            return "Not documented"
        axis_text = f" x {_fmt(axis, 0, ' degrees')}" if axis is not None else " (axis unavailable)"
        return f"{float(sphere):+.2f} D / {float(cylinder):+.2f} D{axis_text}"
    transition = (
        "Not applicable"
        if values.get("transition_zone_mm") is None and values.get("transition_zone_not_applicable") == "yes"
        else _fmt(values.get("transition_zone_mm"), 1, " mm")
    )
    return [
        ("Procedure", _text(values.get("procedure"))),
        ("Prior refractive surgery", _text(values.get("prior_refractive_surgery"))),
        ("Stability / progression / CDVA flag", (
            f"{_text(values.get('refractive_stability'))} / "
            f"{_text(values.get('documented_progression'))} / "
            f"{_text(values.get('unexplained_CDVA_below_20_20'))}"
        )),
        ("Manifest entered notation", entered_refraction("manifest")),
        ("Manifest normalized (minus-cylinder)", manifest),
        ("Intended entered notation", entered_refraction("intended")),
        ("Intended normalized (minus-cylinder)", correction),
        ("Correction source", _text(values.get("correction_source"), "Manual / not documented")),
        ("Score / category", f"{_text(score.get('total'), '-')} / {_text(score.get('category'), '-') }"),
        ("Randleman ERSS / category", f"{_text((eye.get('randleman_erss') or {}).get('total'))} / {_text((eye.get('randleman_erss') or {}).get('category'))}"),
        ("Final BAD-D / class", f"{_fmt((eye.get('bad_summary') or {}).get('value'), 2)} / {_text((eye.get('bad_summary') or {}).get('category'))}"),
        ("CERAI-adapted NICE / class", f"{_text((eye.get('nice') or {}).get('total'))} / {_text((eye.get('nice') or {}).get('category'))}"),
        ("Thinnest pachymetry", _fmt(values.get("pachy_thinnest_um"), 0, " um")),
        ("Manifest MRSE", _fmt(values.get("MRSE_D"), 2, " D")),
        ("Intended MRSE", _fmt(values.get("intended_MRSE_D"), 2, " D")),
        ("Manifest / intended pattern", (
            f"{_text(values.get('manifest_refractive_pattern'))} / "
            f"{_text(values.get('intended_refractive_pattern'))}"
        )),
        ("Intended principal meridians", " / ".join(
            _fmt(item, 2, " D") for item in (values.get("intended_principal_meridians_D") or [])
        ) or "Not documented"),
        ("Preoperative / estimated final Kmean", (
            f"{_fmt(values.get('preoperative_Kmean_D'), 2, ' D')} / "
            f"{_fmt(values.get('estimated_final_Kmean_D'), 2, ' D')}"
        )),
        ("Maximum ablation", _fmt(values.get("max_ablation_um"), 1, " um")),
        ("Laser platform", _text(values.get("laser_platform"))),
        ("PRK epithelium", _fmt(values.get("PRK_epithelium_um"), 0, " um")),
        ("Optical / transition zone", f"{_fmt(values.get('optical_zone_mm'), 1, ' mm')} / {transition}"),
        ("Enhancement anticipated", _text(values.get("enhancement_anticipated"))),
        ("PRK RST / PTA", f"{_fmt(values.get('PRK_RST_um'), 0, ' um')} / {_fmt(values.get('PRK_PTA_percent'), 1, '%')}"),
        ("LASIK RSB / PTA", f"{_fmt(values.get('LASIK_RSB_um'), 0, ' um')} / {_fmt(values.get('LASIK_PTA_percent'), 1, '%')}"),
        ("Tomography review", _text((eye.get("tomography_review") or {}).get("status"))),
        ("Morphology category", _text((eye.get("topography_classification") or {}).get("scoring_category"))),
        ("Randleman I-S / source", (
            f"{_fmt(erss_evidence.get('I_S_D'), 2, ' D')} / "
            f"{_text(erss_evidence.get('I_S_source'))}"
        )),
        ("Validated Randleman topography", (
            f"{_text(erss_evidence.get('validated_category'))} / "
            f"{_text(erss_evidence.get('category_source'))}"
        )),
        ("Pentacam QS", _text(values.get("pentacam_qs"))),
    ]


def _findings(eye: Dict[str, Any]) -> Iterable[tuple[str, List[str]]]:
    bad_display = (eye.get("tomography_review") or {}).get("BAD_display") or {}
    groups = [
        ("Hard stops", eye.get("hard_stops") or []),
        ("Decision reasons", eye.get("reasons") or []),
        ("Missing or unresolved data", eye.get("missing") or []),
        ("Surgical-load evidence flags", eye.get("surgical_load_flags") or []),
        ("Clinical modifiers", eye.get("clinical_modifiers") or []),
        ("Warnings", eye.get("warnings") or []),
        ("NICE component audit", [f"{key}: {value} point(s)" for key, value in (eye.get("nice") or {}).get("rows", {}).items()]
         + [f"Input {key}: {value}" for key, value in (eye.get("nice") or {}).get("values", {}).items()]
         + [f"Source {key}: {value}" for key, value in (eye.get("nice") or {}).get("input_sources", {}).items()]
         + (eye.get("nice") or {}).get("evidence_notes", [])),
        ("NICE interpretation note", [(eye.get("nice") or {})["note"]] if (eye.get("nice") or {}).get("note") else []),
        ("Surgeon attention - hyperopic/mixed pathway", eye.get("surgeon_attention") or []),
        ("Tomography concern flags", (eye.get("tomography_review") or {}).get("cross_sectional_flags") or []),
        ("BAD display interpretation", [f"{key}: {value}" for key, value in bad_display.items()]),
    ]
    return ((title, [str(item) for item in items]) for title, items in groups if items)


def _extracted_eye(extracted: Dict[str, Any], eye_id: str) -> Dict[str, Any]:
    for eye in extracted.get("eyes") or []:
        if eye.get("eye") == eye_id:
            return eye
    return {}


def _tomography_rows(extracted: Dict[str, Any], eye_id: str) -> List[tuple[str, str]]:
    eye = _extracted_eye(extracted, eye_id)
    keys = [
        ("K1", "K1_D", " D", 2), ("K1 axis", "K1_axis_deg", " degrees", 0),
        ("K2", "K2_D", " D", 2), ("K2 axis", "K2_axis_deg", " degrees", 0),
        ("Corneal diameter / W2W", "corneal_diameter_mm", " mm", 2),
        ("Kmax", "Kmax_D", " D", 2), ("Thinnest pachymetry", "pachy_thinnest_um", " um", 0),
        ("BAD-D final", "BAD_D", "", 2), ("BAD-Df", "Df", "", 2),
        ("BAD-Db", "Db", "", 2), ("BAD-Dp", "Dp", "", 2),
        ("BAD-Dt", "Dt", "", 3), ("BAD-Da", "Da", "", 3),
        ("ARTmax", "ARTmax_um", " um", 0), ("PPI min", "PPI_min", "", 2),
        ("PPI avg", "PPI_avg", "", 2), ("PPI max", "PPI_max", "", 2),
        ("ISV", "ISV", "", 0), ("IVA", "IVA", "", 3),
        ("KI", "KI", "", 3), ("CKI", "CKI", "", 3),
        ("IHA", "IHA", "", 3), ("IHD", "IHD", "", 3),
        ("I-S", "I_S", " D", 2), ("KISA", "KISA", "%", 1),
        ("Rmin", "Rmin_mm", " mm", 2), ("SRAX", "srax_deg", " degrees", 1),
        ("Anterior elevation at TP", "anterior_elevation_thinnest_um", " um", 1),
        ("Posterior elevation at TP", "posterior_elevation_thinnest_um", " um", 1),
        ("Thinnest X", "thinnest_x_mm", " mm", 2), ("Thinnest Y", "thinnest_y_mm", " mm", 2),
        ("Corneal volume", "corneal_volume_mm3", " mm3", 2),
        ("RMS-HOA", "RMS_HOA_um", " um", 3), ("Vertical coma", "vertical_coma_um", " um", 3),
        ("Morphology", "morphology", "", 0), ("Anterior pattern", "anterior_pattern", "", 0),
        ("Posterior pattern", "posterior_pattern", "", 0), ("Image quality", "quality", "", 0),
        ("Pentacam QS", "pentacam_qs", "", 0), ("Source files", "source_files", "", 0),
    ]
    return [(label, _fmt(eye.get(key), digits, unit)) for label, key, unit, digits in keys]


def _microkeratome_rows(eye: Dict[str, Any]) -> List[tuple[str, str]]:
    plan = eye.get("microkeratome_planning") or {}
    if not plan.get("applicable"):
        return []
    blades = ", ".join(str(item) for item in plan.get("blade_recommendations") or []) or "Not documented"
    projected = (
        f"{_fmt(plan.get('alternative_rsb_um'), 1, ' um')} / "
        f"{_fmt(plan.get('alternative_pta_percent'), 1, '%')}"
        if plan.get("alternative_safety") != "NOT_APPLICABLE"
        else "Not applicable"
    )
    return [
        ("Assessment gate", _text(plan.get("assessment_gate"))),
        ("Steep-flat K spread", _fmt(plan.get("delta_k_d"), 2, " D")),
        ("Vacuum ring", _fmt(plan.get("vacuum_ring_mm"), 1, " mm")),
        ("Vacuum pressure", _text(plan.get("vacuum_pressure_mmhg"), "Not determined") + " mmHg" if plan.get("vacuum_pressure_mmhg") else "Not determined"),
        ("Blade recommendation(s)", blades),
        ("Primary hinge", _text(plan.get("primary_hinge"), "No CERAI K-spread hinge override")),
        ("Conditional alternative", _text(plan.get("alternative_hinge"), "Not cleared / not applicable")),
        ("Alternative projected RSB / PTA", projected),
        ("Alternative safety", _text(plan.get("alternative_safety"))),
        ("Ring-zone clearance", _fmt(plan.get("ring_tzone_clearance_mm"), 2, " mm")),
        ("Source", _text(plan.get("source"))),
    ]


def _paired_rows(rows: List[tuple[str, str]]) -> List[List[str]]:
    paired: List[List[str]] = []
    for index in range(0, len(rows), 2):
        left = rows[index]
        right = rows[index + 1] if index + 1 < len(rows) else ("", "")
        paired.append([left[0], left[1], right[0], right[1]])
    return paired


def _ordered_eyes(decision: Dict[str, Any]) -> List[Dict[str, Any]]:
    order = {"OD": 0, "OS": 1}
    eyes = [eye for eye in decision.get("eyes") or [] if isinstance(eye, dict)]
    return sorted(eyes, key=lambda eye: order.get(str(eye.get("eye")), 2))


def build_pdf(payload: Dict[str, Any]) -> bytes:
    patient = payload.get("patient") or {}
    decision = payload.get("decision") or {}
    extracted = payload.get("extracted") or {}
    report_eyes = _ordered_eyes(decision)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter, rightMargin=0.65 * inch, leftMargin=0.65 * inch,
        topMargin=0.72 * inch, bottomMargin=0.68 * inch,
        title="CERAI Preoperative Ectasia Risk Assessment",
        author="CERAI",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=17, leading=20, textColor=_rl(NAVY), alignment=TA_LEFT, spaceAfter=3))
    styles.add(ParagraphStyle(name="ReportSub", parent=styles["Normal"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=_rl(GRAY), spaceAfter=12))
    styles.add(ParagraphStyle(name="Section", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=_rl(NAVY), spaceBefore=10, spaceAfter=5))
    styles.add(ParagraphStyle(name="BodySmall", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=_rl(INK), spaceAfter=3))
    styles.add(ParagraphStyle(name="Tiny", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2, leading=9, textColor=_rl(GRAY)))
    styles.add(ParagraphStyle(name="Liability", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9, leading=12, textColor=_rl(RED), backColor=_rl(RED_FILL), borderColor=_rl(RED), borderWidth=0.8, borderPadding=7, spaceAfter=11))
    styles.add(ParagraphStyle(name="PatientName", parent=styles["BodyText"], fontName=PDF_UNICODE_BOLD, fontSize=15, leading=18, textColor=_rl(NAVY), spaceBefore=2, spaceAfter=6))

    story: List[Any] = []
    story.append(Paragraph("CERAI PREOPERATIVE ECTASIA RISK ASSESSMENT", styles["ReportTitle"]))
    story.append(Paragraph(f"Corneal refractive surgery clinical decision-support report | Software v{APP_VERSION}", styles["ReportSub"]))
    story.append(Paragraph(LIABILITY_NOTICE, styles["Liability"]))

    metadata = [
        ["Patient", _ascii(patient.get("name")), "Patient ID", _ascii(patient.get("id"))],
        ["Age", _ascii(patient.get("age")), "Assessment date", _ascii(patient.get("report_date"))],
        ["Reviewer", _ascii(patient.get("reviewer")), "Eyes assessed", ", ".join(_ascii(e.get("eye")) for e in report_eyes) or "None"],
    ]
    meta_table = Table(metadata, colWidths=[0.85 * inch, 2.2 * inch, 1.0 * inch, 2.1 * inch], hAlign="LEFT")
    meta_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"), ("FONTSIZE", (0, 0), (-1, -1), 8.3),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0, 0), (-1, -1), _rl(INK)),
        ("BACKGROUND", (0, 0), (0, -1), _rl(GRAY_FILL)), ("BACKGROUND", (2, 0), (2, -1), _rl(GRAY_FILL)),
        ("GRID", (0, 0), (-1, -1), 0.45, _rl(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 10))
    patient_banner = _ascii(patient.get("name"), "PATIENT NAME NOT DOCUMENTED").upper()
    story.append(Paragraph(patient_banner, styles["PatientName"]))

    overall = _ascii(decision.get("status") or "NOT ASSESSED")
    accent, fill = _status_palette(decision.get("status") or "")
    status_table = Table([[Paragraph(f"<b>OVERALL DISPOSITION</b><br/><font size='13'><b>{overall}</b></font><br/>{_ascii(decision.get('action'), '')}", styles["BodySmall"])]], colWidths=[6.15 * inch])
    status_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _rl(fill)),
        ("BOX", (0, 0), (-1, -1), 1.2, _rl(accent)),
        ("LINEBEFORE", (0, 0), (0, -1), 6, _rl(accent)),
        ("LEFTPADDING", (0, 0), (-1, -1), 12), ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (-1, -1), _rl(accent)),
    ]))
    story.append(status_table)
    identity_warnings = decision.get("identity_warnings") or []
    if identity_warnings:
        story.append(Paragraph("PATIENT IDENTITY NOT VERIFIED - SURGEON CONFIRMATION REQUIRED", styles["Section"]))
        for item in identity_warnings:
            story.append(Paragraph(f"- {_ascii(item)}", styles["BodySmall"]))
    blockers = decision.get("critical_input_issues") or []
    if blockers:
        story.append(Paragraph("Global clinical / source blockers", styles["Section"]))
        for item in blockers:
            story.append(Paragraph(f"- {_ascii(item)}", styles["BodySmall"]))

    for eye in report_eyes:
        eye_status = eye.get("status") or "NOT ASSESSED"
        eye_accent, eye_fill = _status_palette(eye_status)
        story.append(Paragraph(f"{_ascii(eye.get('eye'))} ASSESSMENT", styles["Section"]))
        banner = Table([[_ascii(eye_status), _ascii(eye.get("action"), "")]], colWidths=[1.7 * inch, 4.45 * inch])
        banner.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), _rl(eye_fill)),
            ("TEXTCOLOR", (0, 0), (0, 0), _rl(eye_accent)),
            ("FONTNAME", (0, 0), (0, 0), "Helvetica-Bold"), ("FONTNAME", (1, 0), (1, 0), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("BOX", (0, 0), (-1, -1), 0.6, _rl(eye_accent)),
            ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(banner)
        metric_rows = [["Parameter", "Result"]] + [[_ascii(k), _ascii(v)] for k, v in _eye_metrics(eye)]
        metric_table = Table(metric_rows, colWidths=[2.25 * inch, 3.9 * inch], repeatRows=1)
        metric_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), _rl(NAVY)), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (1, 1), (1, -1), "Helvetica"), ("FONTSIZE", (0, 0), (-1, -1), 7.8),
            ("GRID", (0, 0), (-1, -1), 0.35, _rl(LINE)), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _rl("F7F9FB")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(metric_table)
        for title, items in _findings(eye):
            story.append(Paragraph(_ascii(title), styles["Section"]))
            for item in items:
                story.append(Paragraph(f"- {_ascii(item)}", styles["BodySmall"]))

        planning_rows = _microkeratome_rows(eye)
        if planning_rows:
            story.append(Paragraph("Post-assessment ML7 microkeratome planning", styles["Section"]))
            planning_table = Table(
                [["Parameter", "Surgeon-review recommendation"]]
                + [[_ascii(key), _ascii(value)] for key, value in planning_rows],
                colWidths=[2.25 * inch, 3.9 * inch], repeatRows=1,
            )
            planning_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), _rl(BLUE)), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.8), ("GRID", (0, 0), (-1, -1), 0.35, _rl(LINE)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(planning_table)
            for title, items in (
                ("Planning warnings", (eye.get("microkeratome_planning") or {}).get("warnings") or []),
                ("Planning notes", (eye.get("microkeratome_planning") or {}).get("notes") or []),
            ):
                if items:
                    story.append(Paragraph(title, styles["Section"]))
                    for item in items:
                        story.append(Paragraph(f"- {_ascii(item)}", styles["BodySmall"]))

        story.append(Paragraph("Extracted tomography", styles["Section"]))
        tomo = [["Parameter", "Value", "Parameter", "Value"]] + [
            [_ascii(value, "") for value in row]
            for row in _paired_rows(_tomography_rows(extracted, eye.get("eye")))
        ]
        tomo_table = Table(tomo, colWidths=[1.45 * inch, 1.55 * inch, 1.45 * inch, 1.7 * inch], repeatRows=1)
        tomo_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), _rl(GRAY_FILL)), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"), ("FONTNAME", (2, 1), (2, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7.2),
            ("GRID", (0, 0), (-1, -1), 0.35, _rl(LINE)), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(tomo_table)

    warnings = extracted.get("global_warnings") or []
    if warnings:
        story.append(Paragraph("Extraction warnings", styles["Section"]))
        for item in warnings:
            story.append(Paragraph(f"- {_ascii(item)}", styles["BodySmall"]))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Interpretation note", styles["Section"]))
    story.append(Paragraph(
        "This report is generated under the CERAI Preoperative Ectasia Risk Assessment Protocol for corneal refractive surgery. "
        "CAUTION is a STOP/DEFER decision requiring repeat ectasia/tomographic assessment after at least 6 months. "
        "DATA INSUFFICIENT / NOT ASSESSED does not permit PASS. This clinical decision-support report does not replace independent surgeon review.",
        styles["Tiny"],
    ))

    def page_footer(canvas, pdf_doc):
        canvas.saveState()
        canvas.setStrokeColor(_rl(LINE))
        canvas.line(0.65 * inch, 0.48 * inch, 7.85 * inch, 0.48 * inch)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(_rl(GRAY))
        canvas.drawString(0.65 * inch, 0.32 * inch, "CERAI | Clinical decision-support report")
        canvas.drawRightString(7.85 * inch, 0.32 * inch, f"Page {pdf_doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=page_footer, onLaterPages=page_footer)
    return buffer.getvalue()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _style_doc_table(table, widths: List[float], header=True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
                    run.font.size = Pt(8.5)
    if header:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, NAVY)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def _add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text)


def build_docx(payload: Dict[str, Any]) -> bytes:
    patient = payload.get("patient") or {}
    decision = payload.get("decision") or {}
    extracted = payload.get("extracted") or {}
    report_eyes = _ordered_eyes(decision)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, before, after in ((1, 14, 12, 6), (2, 11, 9, 4)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "CERAI  |  PREOPERATIVE RISK ASSESSMENT"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CERAI | Clinical decision-support report | ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("CERAI PREOPERATIVE ECTASIA RISK ASSESSMENT")
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = document.add_paragraph(f"Corneal refractive surgery clinical decision-support report | Software v{APP_VERSION}")
    subtitle.paragraph_format.space_after = Pt(10)
    for run in subtitle.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    liability = document.add_paragraph()
    liability.paragraph_format.space_after = Pt(10)
    liability.paragraph_format.left_indent = Inches(0.08)
    liability_run = liability.add_run(LIABILITY_NOTICE)
    liability_run.font.name = "Arial"
    liability_run.font.size = Pt(9)
    liability_run.font.bold = True
    liability_run.font.color.rgb = RGBColor.from_string(RED)

    meta = document.add_table(rows=3, cols=4)
    meta.style = "Table Grid"
    rows = [
        ("Patient", _text(patient.get("name")), "Patient ID", _text(patient.get("id"))),
        ("Age", _text(patient.get("age")), "Assessment date", _text(patient.get("report_date"))),
        ("Reviewer", _text(patient.get("reviewer")), "Eyes assessed", ", ".join(_text(e.get("eye")) for e in report_eyes) or "None"),
    ]
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            meta.cell(r_idx, c_idx).text = value
            if c_idx in (0, 2):
                _set_cell_shading(meta.cell(r_idx, c_idx), GRAY_FILL)
                meta.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True
    _style_doc_table(meta, [0.8, 2.1, 1.0, 1.95], header=False)

    patient_banner = document.add_paragraph()
    patient_banner.paragraph_format.space_before = Pt(10)
    patient_banner.paragraph_format.space_after = Pt(6)
    patient_banner_run = patient_banner.add_run(
        _text(patient.get("name"), "PATIENT NAME NOT DOCUMENTED").upper()
    )
    patient_banner_run.font.name = "Arial"
    patient_banner_run.font.size = Pt(15)
    patient_banner_run.font.bold = True
    patient_banner_run.font.color.rgb = RGBColor.from_string(NAVY)

    status = decision.get("status") or "NOT ASSESSED"
    accent, fill = _status_palette(status)
    box = document.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    box.cell(0, 0).text = f"OVERALL DISPOSITION\n{status}\n{_text(decision.get('action'), '')}"
    _set_cell_shading(box.cell(0, 0), fill)
    _set_cell_margins(box.cell(0, 0), top=150, bottom=150, start=180, end=180)
    for idx, run in enumerate(box.cell(0, 0).paragraphs[0].runs):
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor.from_string(accent)
        run.bold = idx < 2

    identity_warnings = decision.get("identity_warnings") or []
    if identity_warnings:
        _add_heading(document, "PATIENT IDENTITY NOT VERIFIED - SURGEON CONFIRMATION REQUIRED", 1)
        for item in identity_warnings:
            _add_bullet(document, str(item))

    blockers = decision.get("critical_input_issues") or []
    if blockers:
        _add_heading(document, "Global clinical / source blockers", 1)
        for item in blockers:
            _add_bullet(document, item)

    for eye in report_eyes:
        _add_heading(document, f"{_text(eye.get('eye'))} assessment", 1)
        eye_status = eye.get("status") or "NOT ASSESSED"
        eye_accent, eye_fill = _status_palette(eye_status)
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        banner.cell(0, 0).text = eye_status
        banner.cell(0, 1).text = _text(eye.get("action"), "")
        for cell in banner.rows[0].cells:
            _set_cell_shading(cell, eye_fill)
            _set_cell_margins(cell, top=110, bottom=110)
        banner.cell(0, 0).paragraphs[0].runs[0].bold = True
        banner.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(eye_accent)
        _style_doc_table(banner, [1.7, 4.15], header=False)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Parameter"
        table.rows[0].cells[1].text = "Result"
        for label, value in _eye_metrics(eye):
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value.replace(" um", " µm")
            cells[0].paragraphs[0].runs[0].bold = True
        _style_doc_table(table, [2.15, 3.7], header=True)

        for heading, items in _findings(eye):
            _add_heading(document, heading, 2)
            for item in items:
                _add_bullet(document, item)

        planning_rows = _microkeratome_rows(eye)
        if planning_rows:
            _add_heading(document, "Post-assessment ML7 microkeratome planning", 2)
            planning_table = document.add_table(rows=1, cols=2)
            planning_table.style = "Table Grid"
            planning_table.rows[0].cells[0].text = "Parameter"
            planning_table.rows[0].cells[1].text = "Surgeon-review recommendation"
            for label, value in planning_rows:
                cells = planning_table.add_row().cells
                cells[0].text = label
                cells[1].text = value.replace(" um", " µm")
                cells[0].paragraphs[0].runs[0].bold = True
            _style_doc_table(planning_table, [2.15, 3.7], header=True)
            for heading, items in (
                ("Planning warnings", (eye.get("microkeratome_planning") or {}).get("warnings") or []),
                ("Planning notes", (eye.get("microkeratome_planning") or {}).get("notes") or []),
            ):
                if items:
                    _add_heading(document, heading, 2)
                    for item in items:
                        _add_bullet(document, str(item))

        _add_heading(document, "Extracted tomography", 2)
        tomo = document.add_table(rows=1, cols=4)
        tomo.style = "Table Grid"
        tomo.rows[0].cells[0].text = "Parameter"
        tomo.rows[0].cells[1].text = "Value"
        tomo.rows[0].cells[2].text = "Parameter"
        tomo.rows[0].cells[3].text = "Value"
        for row in _paired_rows(_tomography_rows(extracted, eye.get("eye"))):
            cells = tomo.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value.replace(" um", " µm")
            cells[0].paragraphs[0].runs[0].bold = True
            if cells[2].paragraphs[0].runs:
                cells[2].paragraphs[0].runs[0].bold = True
        _style_doc_table(tomo, [1.35, 1.575, 1.35, 1.575], header=True)

    warnings = extracted.get("global_warnings") or []
    if warnings:
        _add_heading(document, "Extraction warnings", 1)
        for item in warnings:
            _add_bullet(document, str(item))

    _add_heading(document, "Interpretation note", 1)
    note = document.add_paragraph(
        "This report is generated under the CERAI Preoperative Ectasia Risk Assessment Protocol for corneal refractive surgery. "
        "CAUTION is a STOP/DEFER decision requiring repeat ectasia/tomographic assessment after at least 6 months. "
        "DATA INSUFFICIENT / NOT ASSESSED does not permit PASS. This clinical decision-support report does not replace independent surgeon review."
    )
    note.style = document.styles["Normal"]
    for run in note.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
